#!/usr/bin/env python3
"""
Symbol Extractor
Extracts symbols from a Word document for use as chapter separators.
"""

from docx import Document
from pathlib import Path
import sys


def extract_symbols(doc_path: str, limit: int = 10):
    """Extract symbols from a Word document.
    
    Args:
        doc_path: Path to the Word document containing symbols
        limit: Maximum number of symbols to extract
        
    Returns:
        List of symbols found in the document
    """
    doc = Document(doc_path)
    symbols = []
    
    # First, try to extract from paragraphs
    for para in doc.paragraphs:
        text = para.text.strip()
        if text:
            # Split by whitespace to get individual characters/symbols
            parts = text.split()
            for part in parts:
                # Check various conditions for symbols
                if (len(part) <= 3 and not part.isalnum()) or \
                   (len(part) == 1 and ord(part) > 127) or \
                   any(ord(c) > 8000 for c in part):  # Unicode symbols
                    symbols.append(part)
                    if len(symbols) >= limit * 2:  # Get more initially
                        break
            
            # Also check if the entire paragraph is short and might be symbols
            if len(text) <= 10 and not text.isalnum() and ' ' not in text:
                symbols.append(text)
    
    # Also check tables if they exist
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                text = cell.text.strip()
                if text and len(text) <= 3:
                    # Check if it's a symbol
                    if not text.isalnum() or (len(text) == 1 and ord(text) > 127):
                        symbols.append(text)
                        if len(symbols) >= limit * 2:
                            break
    
    # Remove duplicates while preserving order
    seen = set()
    unique_symbols = []
    for symbol in symbols:
        if symbol not in seen:
            seen.add(symbol)
            unique_symbols.append(symbol)
            if len(unique_symbols) >= limit:
                break
    
    return unique_symbols


def get_first_symbol(doc_path: str) -> str:
    """Get the first symbol from a Word document.
    
    Args:
        doc_path: Path to the Word document
        
    Returns:
        The first symbol found, or a default symbol if none found
    """
    symbols = extract_symbols(doc_path, limit=1)
    return symbols[0] if symbols else "❦"  # Default ornamental symbol


if __name__ == "__main__":
    if len(sys.argv) > 1:
        doc_path = sys.argv[1]
        debug = "--debug" in sys.argv
        
        print(f"Extracting symbols from: {doc_path}")
        
        if debug:
            print("\nDebug mode - showing all paragraphs:")
            doc = Document(doc_path)
            for i, para in enumerate(doc.paragraphs[:20]):  # First 20 paragraphs
                if para.text.strip():
                    print(f"Para {i}: '{para.text.strip()}'")
            
            if doc.tables:
                print("\nTables found:")
                for t_idx, table in enumerate(doc.tables):
                    print(f"\nTable {t_idx + 1}:")
                    for r_idx, row in enumerate(table.rows[:5]):  # First 5 rows
                        row_text = " | ".join([cell.text.strip() for cell in row.cells])
                        if row_text.strip():
                            print(f"  Row {r_idx}: {row_text}")
        
        symbols = extract_symbols(doc_path)
        
        if symbols:
            print(f"\nFound {len(symbols)} symbols:")
            for i, symbol in enumerate(symbols, 1):
                print(f"{i}. {symbol}")
            
            print(f"\nFirst symbol: {symbols[0]}")
        else:
            print("\nNo symbols found in the document.")
            print("Try running with --debug flag to see document content")
    else:
        print("Usage: python symbol_extractor.py <path_to_word_doc> [--debug]")