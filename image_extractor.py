#!/usr/bin/env python3
"""
Image Extractor
Extracts images from Word documents for use as chapter separators.
"""

from docx import Document
from docx.oxml.ns import qn
from pathlib import Path
import io
import sys


def extract_first_image(doc_path: str):
    """Extract the first image from a Word document.
    
    Args:
        doc_path: Path to the Word document containing images
        
    Returns:
        Tuple of (image_data, width, height) or None if no images found
    """
    doc = Document(doc_path)
    
    # Search through all paragraphs for inline images
    for paragraph in doc.paragraphs:
        for run in paragraph.runs:
            # Check if the run contains any inline shapes
            inline_elements = run._element.findall('.//w:drawing', 
                {'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'})
            
            for inline in inline_elements:
                # Find the blip (image reference)
                blip_elements = inline.findall('.//a:blip', 
                    {'a': 'http://schemas.openxmlformats.org/drawingml/2006/main'})
                
                for blip in blip_elements:
                    # Get the relationship ID
                    embed_attr = '{http://schemas.openxmlformats.org/officeDocument/2006/relationships}embed'
                    rId = blip.get(embed_attr)
                    
                    if rId:
                        try:
                            # Get the image part
                            image_part = doc.part.related_parts[rId]
                            image_data = image_part.blob
                            
                            # Try to get image dimensions
                            extent_elements = inline.findall('.//wp:extent',
                                {'wp': 'http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing'})
                            
                            width = None
                            height = None
                            if extent_elements:
                                extent = extent_elements[0]
                                cx = extent.get('cx')
                                cy = extent.get('cy')
                                if cx and cy:
                                    # Convert EMUs to inches
                                    width = int(cx) / 914400
                                    height = int(cy) / 914400
                            
                            return image_data, width, height
                            
                        except Exception as e:
                            print(f"Error extracting image: {e}")
                            continue
    
    # Also check for images in tables
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        inline_elements = run._element.findall('.//w:drawing',
                            {'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'})
                        
                        for inline in inline_elements:
                            blip_elements = inline.findall('.//a:blip',
                                {'a': 'http://schemas.openxmlformats.org/drawingml/2006/main'})
                            
                            for blip in blip_elements:
                                embed_attr = '{http://schemas.openxmlformats.org/officeDocument/2006/relationships}embed'
                                rId = blip.get(embed_attr)
                                
                                if rId:
                                    try:
                                        image_part = doc.part.related_parts[rId]
                                        image_data = image_part.blob
                                        
                                        extent_elements = inline.findall('.//wp:extent',
                                            {'wp': 'http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing'})
                                        
                                        width = None
                                        height = None
                                        if extent_elements:
                                            extent = extent_elements[0]
                                            cx = extent.get('cx')
                                            cy = extent.get('cy')
                                            if cx and cy:
                                                width = int(cx) / 914400
                                                height = int(cy) / 914400
                                        
                                        return image_data, width, height
                                        
                                    except Exception as e:
                                        print(f"Error extracting image from table: {e}")
                                        continue
    
    return None


def save_extracted_image(doc_path: str, output_path: str = None):
    """Extract and save the first image from a Word document.
    
    Args:
        doc_path: Path to the Word document
        output_path: Where to save the image (optional)
    """
    result = extract_first_image(doc_path)
    
    if result:
        image_data, width, height = result
        
        if not output_path:
            # Try to determine file extension from image data
            if image_data[:2] == b'\xff\xd8':
                ext = '.jpg'
            elif image_data[:8] == b'\x89PNG\r\n\x1a\n':
                ext = '.png'
            else:
                ext = '.img'
            output_path = f"extracted_image{ext}"
        
        with open(output_path, 'wb') as f:
            f.write(image_data)
        
        print(f"Image extracted successfully: {output_path}")
        if width and height:
            print(f"Dimensions: {width:.2f} x {height:.2f} inches")
    else:
        print("No images found in the document")


if __name__ == "__main__":
    if len(sys.argv) > 1:
        doc_path = sys.argv[1]
        output_path = sys.argv[2] if len(sys.argv) > 2 else None
        
        print(f"Extracting first image from: {doc_path}")
        save_extracted_image(doc_path, output_path)
    else:
        print("Usage: python image_extractor.py <word_doc> [output_image]")