#!/usr/bin/env python3
"""
Verify that all critical fixes have been applied to the Word document.
This script checks the structure of the generated document.
"""

import sys
from pathlib import Path
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH

def verify_document(doc_path):
    """Verify the document structure and formatting."""
    doc = Document(doc_path)
    results = []
    
    # Track what we find
    found_title = False
    title_has_page_break = False
    found_dedication = False
    dedication_has_page_break = False
    found_contents = False
    contents_has_page_break = False
    found_hierarchical_lists = False
    heading_fonts = []
    
    for i, para in enumerate(doc.paragraphs):
        text = para.text.strip()
        
        # Check Title
        if "The Complete Guide to Spiritual Wisdom" in text:
            found_title = True
            # Check if next paragraph is empty (indicating page break)
            if i + 1 < len(doc.paragraphs) and not doc.paragraphs[i + 1].text.strip():
                title_has_page_break = True
        
        # Check Dedication
        if "Dedicated to" in text:
            found_dedication = True
            # Check for page break after
            if i + 1 < len(doc.paragraphs) and not doc.paragraphs[i + 1].text.strip():
                dedication_has_page_break = True
        
        # Check Contents
        if text == "Contents":
            found_contents = True
            # Check for page break after contents section
            # Look for several paragraphs ahead since contents has items
            for j in range(i + 1, min(i + 10, len(doc.paragraphs))):
                if "Chapter 1" in doc.paragraphs[j].text:
                    # Check if there's empty paragraph before Chapter 1
                    if j > 0 and not doc.paragraphs[j - 1].text.strip():
                        contents_has_page_break = True
                    break
        
        # Check heading fonts
        if para.style and para.style.name and 'Heading' in para.style.name:
            if para.runs:
                font_name = para.runs[0].font.name
                if font_name:
                    heading_fonts.append(f"{para.style.name}: {font_name}")
        
        # Check for hierarchical lists (numbered items)
        if text.startswith("1. The Principle") or text.startswith("1. Presence"):
            found_hierarchical_lists = True
    
    # Print results
    print(f"\nVerification Results for: {doc_path.name}")
    print("=" * 60)
    
    print("\n1. Page Breaks:")
    print(f"   ✅ Title found: {found_title}")
    print(f"   {'✅' if title_has_page_break else '❌'} Title has dedicated page: {title_has_page_break}")
    print(f"   ✅ Dedication found: {found_dedication}")
    print(f"   {'✅' if dedication_has_page_break else '❌'} Dedication has dedicated page: {dedication_has_page_break}")
    print(f"   ✅ Contents found: {found_contents}")
    print(f"   {'✅' if contents_has_page_break else '❌'} Contents has dedicated page: {contents_has_page_break}")
    
    print("\n2. Heading Fonts:")
    if heading_fonts:
        for hf in heading_fonts[:5]:  # Show first 5
            print(f"   {hf}")
        # Check if any use Calibri
        calibri_count = sum(1 for hf in heading_fonts if 'Calibri' in hf)
        if calibri_count > 0:
            print(f"   ❌ Found {calibri_count} headings using Calibri font")
        else:
            print(f"   ✅ No headings using Calibri font")
    else:
        print("   ⚠️  No headings found")
    
    print("\n3. Hierarchical Lists:")
    print(f"   {'✅' if found_hierarchical_lists else '❌'} Hierarchical lists found: {found_hierarchical_lists}")
    
    print("\n" + "=" * 60)
    
    # Overall assessment
    all_page_breaks_ok = title_has_page_break and dedication_has_page_break and contents_has_page_break
    no_calibri = all('Calibri' not in hf for hf in heading_fonts)
    
    if all_page_breaks_ok and no_calibri and found_hierarchical_lists:
        print("✅ All critical fixes appear to be working correctly!")
    else:
        print("⚠️  Some issues may still need attention")
    
    return {
        'page_breaks_ok': all_page_breaks_ok,
        'fonts_ok': no_calibri,
        'lists_ok': found_hierarchical_lists
    }

if __name__ == "__main__":
    if len(sys.argv) > 1:
        doc_path = Path(sys.argv[1])
    else:
        # Default to test file
        doc_path = Path("test_samples/test_all_fixes_formatted.docx")
    
    if doc_path.exists():
        verify_document(doc_path)
    else:
        print(f"Error: Document not found: {doc_path}")
        sys.exit(1)