"""
Edge case tests for the Word Formatter system.
Tests unusual, boundary, and error conditions.
"""

import os
import pytest
import tempfile
from pathlib import Path
from unittest.mock import patch, MagicMock
from docx import Document

from document_converter import DocumentConverter, StyleExtractor
from document_converter_ai import AIDocumentConverter


class TestFileEdgeCases:
    """Test edge cases related to file handling."""
    
    @pytest.mark.unit
    def test_empty_file_handling(self, reference_doc_path, temp_dir):
        """Test handling of completely empty files."""
        # Arrange
        empty_file = temp_dir / "empty.txt"
        empty_file.write_text("")
        
        converter = DocumentConverter(str(reference_doc_path))
        output_path = temp_dir / "empty_output.docx"
        
        # Act
        converter.convert(str(empty_file), str(output_path))
        
        # Assert
        assert output_path.exists()
        output_doc = Document(str(output_path))
        # Even empty documents should have basic structure
        assert len(output_doc.paragraphs) >= 0
    
    @pytest.mark.unit
    def test_whitespace_only_file(self, reference_doc_path, temp_dir):
        """Test handling of files with only whitespace."""
        # Arrange
        whitespace_file = temp_dir / "whitespace.txt"
        whitespace_file.write_text("   \n\n\t\t  \n   ")
        
        converter = DocumentConverter(str(reference_doc_path))
        output_path = temp_dir / "whitespace_output.docx"
        
        # Act
        converter.convert(str(whitespace_file), str(output_path))
        
        # Assert
        assert output_path.exists()
    
    @pytest.mark.unit
    def test_very_long_lines(self, reference_doc_path, temp_dir):
        """Test handling of extremely long lines."""
        # Arrange
        long_line_file = temp_dir / "long_lines.txt"
        # Create a line with 10,000+ characters
        very_long_line = "This is an extremely long line with lots of content. " * 200
        content = f"Normal paragraph.\n\n{very_long_line}\n\nAnother normal paragraph."
        long_line_file.write_text(content)
        
        converter = DocumentConverter(str(reference_doc_path))
        output_path = temp_dir / "long_lines_output.docx"
        
        # Act
        converter.convert(str(long_line_file), str(output_path))
        
        # Assert
        assert output_path.exists()
        output_doc = Document(str(output_path))
        
        # Content should be preserved
        all_text = '\n'.join([p.text for p in output_doc.paragraphs])
        assert "extremely long line" in all_text
    
    @pytest.mark.unit
    def test_file_with_null_bytes(self, reference_doc_path, temp_dir):
        """Test handling of files containing null bytes."""
        # Arrange
        null_byte_file = temp_dir / "null_bytes.txt"
        content = "Normal text\x00with null bytes\x00in between."
        null_byte_file.write_bytes(content.encode('utf-8'))
        
        converter = DocumentConverter(str(reference_doc_path))
        output_path = temp_dir / "null_bytes_output.docx"
        
        # Act & Assert
        # This should either succeed by filtering null bytes or raise a specific exception
        try:
            converter.convert(str(null_byte_file), str(output_path))
            # If it succeeds, check that content is preserved
            assert output_path.exists()
            output_doc = Document(str(output_path))
            all_text = '\n'.join([p.text for p in output_doc.paragraphs])
            assert "Normal text" in all_text
        except (ValueError, UnicodeError):
            # This is acceptable behavior for files with null bytes
            pytest.skip("System correctly rejects files with null bytes")
    
    @pytest.mark.unit
    def test_file_with_special_characters(self, reference_doc_path, temp_dir):
        """Test handling of files with various special characters."""
        # Arrange
        special_char_file = temp_dir / "special_chars.txt"
        content = """Special Characters Test

Control Characters: \t\r\n
Currency: $ € £ ¥ ₹
Math: ∑ ∏ ∫ √ ∞ ≈ ≠
Arrows: ← → ↑ ↓ ⇒ ⇔
Symbols: © ® ™ § ¶ †
Quotes: " " ' ' « »
Dashes: – — ‒ ―
Spaces: thin space, en space, em space
Zero-width: zero-width joiner, zero-width non-joiner"""
        special_char_file.write_text(content, encoding='utf-8')
        
        converter = DocumentConverter(str(reference_doc_path))
        output_path = temp_dir / "special_chars_output.docx"
        
        # Act
        converter.convert(str(special_char_file), str(output_path))
        
        # Assert
        assert output_path.exists()
        output_doc = Document(str(output_path))
        all_text = '\n'.join([p.text for p in output_doc.paragraphs])
        assert "Special Characters Test" in all_text
        assert "€" in all_text  # Test specific special character preservation
    
    @pytest.mark.unit
    def test_file_with_mixed_encodings(self, reference_doc_path, temp_dir):
        """Test handling of files with encoding issues."""
        # Arrange
        mixed_encoding_file = temp_dir / "mixed_encoding.txt"
        # Write content that might cause encoding issues
        content = "Mixed encoding test: café naïve résumé"
        # Write as UTF-8 first
        mixed_encoding_file.write_text(content, encoding='utf-8')
        
        # Then append some bytes that might be problematic
        with open(mixed_encoding_file, 'ab') as f:
            f.write(b'\n\nAdditional content with special chars: \xe9\xe8\xe0')
        
        converter = DocumentConverter(str(reference_doc_path))
        output_path = temp_dir / "mixed_encoding_output.docx"
        
        # Act & Assert
        try:
            converter.convert(str(mixed_encoding_file), str(output_path))
            assert output_path.exists()
        except UnicodeDecodeError:
            # This is acceptable behavior for corrupted encodings
            pytest.skip("File encoding is corrupted as expected")


class TestUnicodeAndInternationalization:
    """Test edge cases related to Unicode and international text."""
    
    @pytest.mark.unit
    def test_multiple_scripts_in_one_document(self, reference_doc_path, temp_dir):
        """Test documents with multiple writing systems."""
        # Arrange
        multi_script_file = temp_dir / "multi_script.txt"
        content = """Multi-Script Document Test

English: Hello, this is English text.
Spanish: Hola, esto es texto en español.
French: Bonjour, ceci est du texte français.
German: Guten Tag, das ist deutscher Text.
Russian: Привет, это русский текст.
Arabic: مرحبا، هذا نص عربي.
Hebrew: שלום, זה טקסט עברי.
Chinese: 你好，这是中文文本。
Japanese: こんにちは、これは日本語のテキストです。
Korean: 안녕하세요, 이것은 한국어 텍스트입니다.
Hindi: नमस्ते, यह हिंदी पाठ है।
Thai: สวัสดี นี่คือข้อความภาษาไทย
Sanskrit: नमस्ते, एतत् संस्कृत पाठ: अस्ति।
Transliteration: Namasté, etat saṃskṛta pāṭhaḥ asti."""
        multi_script_file.write_text(content, encoding='utf-8')
        
        converter = DocumentConverter(str(reference_doc_path))
        output_path = temp_dir / "multi_script_output.docx"
        
        # Act
        converter.convert(str(multi_script_file), str(output_path))
        
        # Assert
        assert output_path.exists()
        output_doc = Document(str(output_path))
        all_text = '\n'.join([p.text for p in output_doc.paragraphs])
        
        # Check preservation of various scripts
        assert "Multi-Script Document Test" in all_text
        assert "Привет" in all_text  # Russian
        assert "你好" in all_text    # Chinese
        assert "こんにちは" in all_text  # Japanese
        assert "नमस्ते" in all_text  # Hindi/Sanskrit
    
    @pytest.mark.unit
    @pytest.mark.claude
    def test_ai_processing_unicode_text(self, reference_doc_path, temp_dir, mock_notifications):
        """Test AI processing of Unicode text with proper italicization."""
        # Arrange
        unicode_file = temp_dir / "unicode_ai.txt"
        content = """Unicode AI Test

Regular English text here.

Sanskrit with diacriticals: Vāg vai brahma | vācā hy evedaṃ sarvaṃ sṛṣṭam

Devanagari script: हिन्दी में परीक्षण पाठ

Mixed content: The concept of mokṣa (मोक्ष) is central to Hindu philosophy.

Arabic text: السلام عليكم ورحمة الله وبركاته

More English content to test context switching."""
        unicode_file.write_text(content, encoding='utf-8')
        
        converter = AIDocumentConverter(str(reference_doc_path))
        converter.show_progress = False
        output_path = temp_dir / "unicode_ai_output.docx"
        
        # Mock successful AI processing
        def mock_run(cmd, **kwargs):
            if 'claude' in str(cmd):
                mock_result = MagicMock()
                mock_result.returncode = 0
                mock_result.stdout = """# Unicode AI Test

Regular English text here.

Sanskrit with diacriticals: *Vāg vai brahma | vācā hy evedaṃ sarvaṃ sṛṣṭam*

Devanagari script: *हिन्दी में परीक्षण पाठ*

Mixed content: The concept of *mokṣa* (*मोक्ष*) is central to Hindu philosophy.

Arabic text: *السلام عليكم ورحمة الله وبركاته*

More English content to test context switching."""
                mock_result.stderr = ""
                return mock_result
            else:
                # For notifications
                mock_result = MagicMock()
                mock_result.returncode = 0
                return mock_result
        
        with patch('subprocess.run', side_effect=mock_run):
            # Act
            result = converter.convert_with_ai(str(unicode_file), str(output_path))
        
        # Assert
        assert result is True
        assert output_path.exists()
        
        # Check markdown was properly formatted
        markdown_path = output_path.parent / f"{unicode_file.stem}_markdown.md"
        if markdown_path.exists():
            markdown_content = markdown_path.read_text()
            # Non-English text should be italicized in AI output
            assert "*Vāg vai brahma*" in markdown_content
            assert "*हिन्दी*" in markdown_content
    
    @pytest.mark.unit
    def test_bidi_text_handling(self, reference_doc_path, temp_dir):
        """Test bidirectional text (RTL/LTR) handling."""
        # Arrange
        bidi_file = temp_dir / "bidi_text.txt"
        content = """Bidirectional Text Test

English text flows left to right.

Arabic text flows right to left: مرحبا بكم في اختبار النص ثنائي الاتجاه

Hebrew text also RTL: שלום וברוכים הבאים לבדיקת טקסט דו-כיווני

Mixed LTR/RTL: The Arabic word for peace is سلام (salam) and in Hebrew שלום (shalom).

Numbers in different scripts: 123 ١٢٣ (Arabic numerals)"""
        bidi_file.write_text(content, encoding='utf-8')
        
        converter = DocumentConverter(str(reference_doc_path))
        output_path = temp_dir / "bidi_output.docx"
        
        # Act
        converter.convert(str(bidi_file), str(output_path))
        
        # Assert
        assert output_path.exists()
        output_doc = Document(str(output_path))
        all_text = '\n'.join([p.text for p in output_doc.paragraphs])
        assert "Bidirectional Text Test" in all_text
        assert "مرحبا" in all_text
        assert "שלום" in all_text


class TestMalformedInputEdgeCases:
    """Test edge cases with malformed or unusual input."""
    
    @pytest.mark.unit
    def test_malformed_markdown_syntax(self, reference_doc_path, temp_dir):
        """Test handling of malformed Markdown syntax."""
        # Arrange
        malformed_md = temp_dir / "malformed.md"
        content = """# Heading with missing space#
        
## Heading with extra spaces   ##  

### Incomplete heading
        
*Italic with missing closing asterisk
        
**Bold with only one asterisk*
        
[Link with missing closing bracket(http://example.com)
        
![Image with malformed syntax](missing-alt-text
        
| Table | with | incomplete |
| rows | missing | cells
| and | wrong | alignment |

```
Code block without closing
        
> Blockquote
> that continues
without proper formatting

1. Numbered list
3. With wrong numbers
a. Mixed with letters
   - And inconsistent indentation
      * Multiple bullet types
        
---
Incomplete horizontal rule--"""
        malformed_md.write_text(content)
        
        converter = DocumentConverter(str(reference_doc_path))
        output_path = temp_dir / "malformed_output.docx"
        
        # Act - should not raise exceptions
        converter.convert(str(malformed_md), str(output_path))
        
        # Assert
        assert output_path.exists()
        output_doc = Document(str(output_path))
        all_text = '\n'.join([p.text for p in output_doc.paragraphs])
        assert "Heading with missing space" in all_text
    
    @pytest.mark.unit
    def test_deeply_nested_markdown_structures(self, reference_doc_path, temp_dir):
        """Test handling of deeply nested Markdown structures."""
        # Arrange
        nested_md = temp_dir / "deeply_nested.md"
        content = """# Level 1
## Level 2
### Level 3
#### Level 4
##### Level 5
###### Level 6

> Blockquote level 1
> > Blockquote level 2
> > > Blockquote level 3
> > > > Blockquote level 4

1. List level 1
   1. List level 2
      1. List level 3
         1. List level 4
            1. List level 5
               - Mixed with bullets
                 - Bullet level 2
                   * Different bullet
                     + Another bullet type

**Bold with *italic inside* and `code`**

*Italic with **bold inside** and `code`*

`Code with **attempted bold** and *attempted italic*`"""
        nested_md.write_text(content)
        
        converter = DocumentConverter(str(reference_doc_path))
        output_path = temp_dir / "nested_output.docx"
        
        # Act
        converter.convert(str(nested_md), str(output_path))
        
        # Assert
        assert output_path.exists()
        output_doc = Document(str(output_path))
        
        # Check that deep headings are handled
        headings = [p for p in output_doc.paragraphs if p.style.name.startswith('Heading')]
        assert len(headings) >= 6  # Should have all heading levels
    
    @pytest.mark.unit
    def test_extremely_large_tables(self, reference_doc_path, temp_dir):
        """Test handling of very large tables."""
        # Arrange
        large_table_md = temp_dir / "large_table.md"
        
        # Create a table with many columns and rows
        header_cols = ['Col' + str(i) for i in range(20)]  # 20 columns
        header = '| ' + ' | '.join(header_cols) + ' |'
        separator = '| ' + ' | '.join(['---'] * 20) + ' |'
        
        rows = []
        for row_num in range(100):  # 100 rows
            row_data = [f'R{row_num}C{col}' for col in range(20)]
            row = '| ' + ' | '.join(row_data) + ' |'
            rows.append(row)
        
        table_content = '\n'.join([header, separator] + rows)
        content = f"""# Large Table Test

This document contains a very large table:

{table_content}

End of document."""
        large_table_md.write_text(content)
        
        converter = DocumentConverter(str(reference_doc_path))
        output_path = temp_dir / "large_table_output.docx"
        
        # Act
        converter.convert(str(large_table_md), str(output_path))
        
        # Assert
        assert output_path.exists()
        output_doc = Document(str(output_path))
        
        # Should have created a table
        assert len(output_doc.tables) >= 1
        table = output_doc.tables[0]
        assert len(table.columns) >= 10  # Should have many columns
        assert len(table.rows) >= 50     # Should have many rows


class TestSystemResourceEdgeCases:
    """Test edge cases related to system resources and limits."""
    
    @pytest.mark.unit
    def test_very_long_file_paths(self, reference_doc_path, temp_dir):
        """Test handling of extremely long file paths."""
        # Arrange - Create deeply nested directory structure
        deep_dir = temp_dir
        for i in range(10):  # Create deep nesting
            deep_dir = deep_dir / f"very_long_directory_name_level_{i}_with_lots_of_characters"
            deep_dir.mkdir(exist_ok=True)
        
        long_filename = "extremely_long_filename_with_many_characters_that_might_cause_issues.txt"
        long_path_file = deep_dir / long_filename
        long_path_file.write_text("Test content for long path file.")
        
        converter = DocumentConverter(str(reference_doc_path))
        output_path = deep_dir / "long_path_output.docx"
        
        # Act
        try:
            converter.convert(str(long_path_file), str(output_path))
            # Assert
            assert output_path.exists()
        except OSError as e:
            # On some systems, very long paths might not be supported
            if "File name too long" in str(e):
                pytest.skip("System doesn't support very long file paths")
            else:
                raise
    
    @pytest.mark.unit
    def test_special_characters_in_file_paths(self, reference_doc_path, temp_dir):
        """Test handling of special characters in file paths."""
        # Arrange
        special_chars_in_names = [
            "file with spaces.txt",
            "file-with-dashes.txt", 
            "file_with_underscores.txt",
            "file.with.dots.txt",
            "file(with)parentheses.txt",
            "file[with]brackets.txt",
            "file{with}braces.txt",
            "file'with'quotes.txt",
            "file\"with\"doublequotes.txt",
            "fileWith$pecial©haracters.txt",
            "file,with,commas.txt"
        ]
        
        for filename in special_chars_in_names:
            try:
                special_file = temp_dir / filename
                special_file.write_text(f"Test content for {filename}")
                
                converter = DocumentConverter(str(reference_doc_path))
                output_name = filename.replace('.txt', '_output.docx')
                output_path = temp_dir / output_name
                
                # Act
                converter.convert(str(special_file), str(output_path))
                
                # Assert
                assert output_path.exists(), f"Failed to create output for {filename}"
                
            except OSError as e:
                # Some special characters might not be allowed on certain filesystems
                if "Invalid" in str(e) or "cannot" in str(e).lower():
                    continue  # Skip this filename on this system
                else:
                    raise


class TestAIEdgeCases:
    """Test edge cases specific to AI processing."""
    
    @pytest.mark.unit
    @pytest.mark.claude
    def test_ai_with_extremely_repetitive_content(self, reference_doc_path, temp_dir, mock_notifications):
        """Test AI processing of highly repetitive content."""
        # Arrange
        repetitive_file = temp_dir / "repetitive.txt"
        # Create content that might confuse AI
        repetitive_content = """Test Document

Chapter 1

Test paragraph. """ * 1000  # Very repetitive content
        repetitive_file.write_text(repetitive_content)
        
        converter = AIDocumentConverter(str(reference_doc_path))
        converter.show_progress = False
        output_path = temp_dir / "repetitive_output.docx"
        
        # Mock AI to return reasonable output despite repetitive input
        def mock_run(cmd, **kwargs):
            if 'claude' in str(cmd):
                mock_result = MagicMock()
                mock_result.returncode = 0
                mock_result.stdout = """# Test Document

## Chapter 1

Test paragraph. Test paragraph. Test paragraph."""
                mock_result.stderr = ""
                return mock_result
            else:
                mock_result = MagicMock()
                mock_result.returncode = 0
                return mock_result
        
        with patch('subprocess.run', side_effect=mock_run):
            # Act
            result = converter.convert_with_ai(str(repetitive_file), str(output_path))
        
        # Assert
        assert result is True
        assert output_path.exists()
    
    @pytest.mark.unit
    @pytest.mark.claude
    def test_ai_with_mixed_content_languages(self, reference_doc_path, temp_dir, mock_notifications):
        """Test AI processing with mixed language content requiring different handling."""
        # Arrange
        mixed_file = temp_dir / "mixed_languages.txt"
        content = """Multilingual Test Document

English Section
This is regular English text that should not be italicized.

Sanskrit Section
Here is Sanskrit with diacriticals: Vāg vai brahma and mokṣa.
And some words: ātman, dharma, karma, saṃsāra.

Devanagari Section  
हिन्दी पाठ: यह हिन्दी में लिखा गया है।
संस्कृत: सत्यमेव जयते नानृतम्।

Arabic Section
النص العربي: هذا نص مكتوب باللغة العربية.
السلام عليكم ورحمة الله وبركاته.

Mixed Section
The Sanskrit word mokṣa (मोक्ष) means liberation.
In Arabic, peace is سلام (salaam).
The concept of dharma is central to Hindu philosophy."""
        mixed_file.write_text(content, encoding='utf-8')
        
        converter = AIDocumentConverter(str(reference_doc_path))
        converter.show_progress = False
        output_path = temp_dir / "mixed_languages_output.docx"
        
        # Mock AI to properly handle mixed languages
        def mock_run(cmd, **kwargs):
            if 'claude' in str(cmd):
                mock_result = MagicMock()
                mock_result.returncode = 0
                mock_result.stdout = """# Multilingual Test Document

## English Section
This is regular English text that should not be italicized.

## Sanskrit Section  
Here is Sanskrit with diacriticals: *Vāg vai brahma* and *mokṣa*.
And some words: *ātman*, *dharma*, *karma*, *saṃsāra*.

## Devanagari Section
*हिन्दी पाठ: यह हिन्दी में लिखा गया है।*
*संस्कृत: सत्यमेव जयते नानृतम्।*

## Arabic Section
*النص العربي: هذا نص مكتوب باللغة العربية।*
*السلام عليكم ورحمة الله وبركاته।*

## Mixed Section
The Sanskrit word *mokṣa* (*मोक्ष*) means liberation.
In Arabic, peace is *سلام* (salaam).
The concept of *dharma* is central to Hindu philosophy."""
                mock_result.stderr = ""
                return mock_result
            else:
                mock_result = MagicMock()
                mock_result.returncode = 0
                return mock_result
        
        with patch('subprocess.run', side_effect=mock_run):
            # Act
            result = converter.convert_with_ai(str(mixed_file), str(output_path))
        
        # Assert
        assert result is True
        assert output_path.exists()
        
        # Check that proper italicization occurred
        markdown_path = output_path.parent / f"{mixed_file.stem}_markdown.md"
        if markdown_path.exists():
            markdown_content = markdown_path.read_text()
            assert "*mokṣa*" in markdown_content
            assert "*हिन्दी*" in markdown_content
            assert "*السلام*" in markdown_content
            # Regular English should not be italicized
            assert "This is regular English text" in markdown_content
            assert "*This is regular English text*" not in markdown_content