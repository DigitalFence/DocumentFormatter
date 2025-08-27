"""
Pytest configuration and shared fixtures for the Word Formatter test suite.
"""

import os
import sys
import tempfile
import shutil
from pathlib import Path
from unittest.mock import MagicMock, patch
from typing import Generator, Dict, Any

import pytest
from docx import Document
from docx.shared import Pt, RGBColor


# Add the parent directory to sys.path so we can import the modules being tested
TEST_DIR = Path(__file__).parent
PROJECT_ROOT = TEST_DIR.parent
sys.path.insert(0, str(PROJECT_ROOT))

# Import modules after adding to path
from document_converter import DocumentConverter, StyleExtractor
from document_converter_ai import AIDocumentConverter


@pytest.fixture(scope="session")
def test_data_dir() -> Path:
    """Return the path to the test data directory."""
    return TEST_DIR / "fixtures" / "sample_docs"


@pytest.fixture(scope="session")
def reference_doc_path(test_data_dir: Path) -> Path:
    """Create and return a test reference document."""
    ref_path = test_data_dir / "test_reference.docx"
    if ref_path.exists():
        return ref_path
    
    # Create a simple reference document for testing
    doc = Document()
    
    # Add a title
    title = doc.add_heading("Sample Title", level=0)
    title.runs[0].font.name = "Times New Roman"
    title.runs[0].font.size = Pt(16)
    title.runs[0].font.bold = True
    
    # Add headings
    h1 = doc.add_heading("Chapter 1", level=1)
    h1.runs[0].font.name = "Times New Roman"
    h1.runs[0].font.size = Pt(14)
    
    h2 = doc.add_heading("Section 1.1", level=2)
    h2.runs[0].font.name = "Times New Roman"
    h2.runs[0].font.size = Pt(12)
    
    # Add normal paragraph
    para = doc.add_paragraph("This is a sample paragraph with normal text.")
    for run in para.runs:
        run.font.name = "Times New Roman"
        run.font.size = Pt(11)
    
    # Add another paragraph with bold/italic
    para2 = doc.add_paragraph()
    run1 = para2.add_run("This is ")
    run1.font.name = "Times New Roman"
    run1.font.size = Pt(11)
    
    run2 = para2.add_run("bold text")
    run2.font.name = "Times New Roman"
    run2.font.size = Pt(11)
    run2.bold = True
    
    run3 = para2.add_run(" and ")
    run3.font.name = "Times New Roman"
    run3.font.size = Pt(11)
    
    run4 = para2.add_run("italic text")
    run4.font.name = "Times New Roman"
    run4.font.size = Pt(11)
    run4.italic = True
    
    run5 = para2.add_run(".")
    run5.font.name = "Times New Roman"
    run5.font.size = Pt(11)
    
    # Set margins
    section = doc.sections[0]
    section.top_margin = Pt(72)    # 1 inch
    section.bottom_margin = Pt(72) # 1 inch
    section.left_margin = Pt(72)   # 1 inch
    section.right_margin = Pt(72)  # 1 inch
    
    # Create directory if it doesn't exist
    test_data_dir.mkdir(parents=True, exist_ok=True)
    doc.save(str(ref_path))
    return ref_path


@pytest.fixture
def temp_dir() -> Generator[Path, None, None]:
    """Create a temporary directory for test files."""
    with tempfile.TemporaryDirectory() as temp_dir:
        yield Path(temp_dir)


@pytest.fixture
def sample_text_file(temp_dir: Path) -> Path:
    """Create a sample text file for testing."""
    text_file = temp_dir / "sample.txt"
    content = """Test Document

Chapter 1: Introduction

This is the first paragraph of our test document. It should be converted properly.

This is a second paragraph with some special content like Vāg vai brahma and हिन्दी text.

Section 1.1: Subsection

Here is a bulleted list:
- First item
- Second item
- Third item

And here's a numbered list:
1. First numbered item
2. Second numbered item
3. Third numbered item

Chapter 2: Advanced Topics

More content here with some code snippets and technical terms.
"""
    text_file.write_text(content, encoding='utf-8')
    return text_file


@pytest.fixture
def sample_markdown_file(temp_dir: Path) -> Path:
    """Create a sample markdown file for testing."""
    md_file = temp_dir / "sample.md"
    content = """# Test Document

## Chapter 1: Introduction

This is the first paragraph of our test document. It should be converted properly.

This is a second paragraph with some *special content* like *Vāg vai brahma* and *हिन्दी* text.

### Section 1.1: Subsection

Here is a bulleted list:
- First item
- Second item
- Third item

And here's a numbered list:
1. First numbered item
2. Second numbered item
3. Third numbered item

## Chapter 2: Advanced Topics

More content here with some `code snippets` and **technical terms**.

| Column 1 | Column 2 | Column 3 |
|----------|----------|----------|
| Data 1   | Data 2   | Data 3   |
| Data 4   | Data 5   | Data 6   |
"""
    md_file.write_text(content, encoding='utf-8')
    return md_file


@pytest.fixture
def sample_rtf_file(temp_dir: Path) -> Path:
    """Create a sample RTF file for testing."""
    rtf_file = temp_dir / "sample.rtf"
    # Simple RTF content
    rtf_content = r"""{\rtf1\ansi\deff0 {\fonttbl {\f0 Times New Roman;}}
\f0\fs24 Test Document\par
\par
This is a sample RTF document with some text.\par
\par
It has multiple paragraphs and should be processed correctly.\par
}"""
    rtf_file.write_text(rtf_content, encoding='utf-8')
    return rtf_file


@pytest.fixture
def empty_text_file(temp_dir: Path) -> Path:
    """Create an empty text file for testing edge cases."""
    empty_file = temp_dir / "empty.txt"
    empty_file.write_text("", encoding='utf-8')
    return empty_file


@pytest.fixture
def large_text_file(temp_dir: Path) -> Path:
    """Create a large text file for performance testing."""
    large_file = temp_dir / "large.txt"
    content = "This is a test paragraph with some content. " * 500  # ~20KB per chapter
    chapters = []
    for i in range(10):  # Create ~200KB file
        chapters.append(f"Chapter {i+1}: Test Chapter\n\n{content}\n\n")
    
    large_file.write_text('\n'.join(chapters), encoding='utf-8')
    return large_file


@pytest.fixture
def sanskrit_text_file(temp_dir: Path) -> Path:
    """Create a text file with Sanskrit and non-English content."""
    sanskrit_file = temp_dir / "sanskrit.txt"
    content = """Test Document with Sanskrit

Introduction

This document contains Sanskrit transliteration: Vāg vai brahma | vācā hy evedaṃ sarvaṃ sṛṣṭam

Devanagari Script: हिन्दी में परीक्षण

Arabic: السلام عليكم

Chinese: 你好世界

Regular English text should not be italicized.
"""
    sanskrit_file.write_text(content, encoding='utf-8')
    return sanskrit_file


@pytest.fixture
def mock_claude_success():
    """Mock successful Claude CLI responses."""
    def mock_run(cmd, **kwargs):
        mock_result = MagicMock()
        mock_result.returncode = 0
        mock_result.stdout = """# Test Document

## Chapter 1: Introduction

This is the first paragraph of our test document. It should be converted properly.

This is a second paragraph with some *special content* like *Vāg vai brahma* and *हिन्दी* text.

### Section 1.1: Subsection

Here is a bulleted list:
- First item
- Second item  
- Third item

## Chapter 2: Advanced Topics

More content here with technical terms.
"""
        mock_result.stderr = ""
        return mock_result
    
    with patch('subprocess.run', side_effect=mock_run):
        yield


@pytest.fixture
def mock_claude_timeout():
    """Mock Claude CLI timeout."""
    def mock_run(cmd, **kwargs):
        import subprocess
        raise subprocess.TimeoutExpired(cmd, kwargs.get('timeout', 600))
    
    with patch('subprocess.run', side_effect=mock_run):
        yield


@pytest.fixture
def mock_claude_error():
    """Mock Claude CLI error responses."""
    def mock_run(cmd, **kwargs):
        mock_result = MagicMock()
        mock_result.returncode = 1
        mock_result.stdout = ""
        mock_result.stderr = "Claude API error: Rate limit exceeded"
        return mock_result
    
    with patch('subprocess.run', side_effect=mock_run):
        yield


@pytest.fixture
def mock_claude_incomplete():
    """Mock Claude returning incomplete response."""
    def mock_run(cmd, **kwargs):
        mock_result = MagicMock()
        mock_result.returncode = 0
        mock_result.stdout = """# Test Document

## Chapter 1: Introduction

This is the first paragraph of our test document.

Would you like me to continue with the rest of the document?"""
        mock_result.stderr = ""
        return mock_result
    
    with patch('subprocess.run', side_effect=mock_run):
        yield


@pytest.fixture
def mock_notifications():
    """Mock system notifications to prevent actual notifications during tests."""
    with patch('subprocess.run') as mock_run:
        # Only mock notification calls, let other subprocess calls through
        def selective_mock(cmd, **kwargs):
            if isinstance(cmd, list) and len(cmd) > 0:
                if 'terminal-notifier' in cmd[0] or 'osascript' in cmd[0]:
                    mock_result = MagicMock()
                    mock_result.returncode = 0
                    return mock_result
            # For non-notification calls, use the real subprocess.run
            import subprocess
            return subprocess.run(cmd, **kwargs)
        
        mock_run.side_effect = selective_mock
        yield mock_run


@pytest.fixture(autouse=True)
def clean_environment():
    """Clean environment variables before each test."""
    env_vars_to_clean = [
        'WORD_FORMATTER_DEBUG',
        'SAVE_MARKDOWN',
        'SHOW_PROGRESS',
        'CLAUDE_MODEL',
        'CLAUDE_TIMEOUT',
        'ENABLE_HAIKU_FALLBACK',
        'CHUNK_THRESHOLD',
        'CLAUDE_CLI_PATH',
        'ERROR_LOG'
    ]
    
    # Store original values
    original_values = {}
    for var in env_vars_to_clean:
        original_values[var] = os.environ.get(var)
        if var in os.environ:
            del os.environ[var]
    
    yield
    
    # Restore original values
    for var, value in original_values.items():
        if value is not None:
            os.environ[var] = value
        elif var in os.environ:
            del os.environ[var]


# Test data for parameterized tests
TEST_FILE_TYPES = [
    ("txt", "sample.txt"),
    ("md", "sample.md"),
    ("rtf", "sample.rtf")
]

CLAUDE_MODELS = ["sonnet", "haiku", "opus"]

ERROR_SCENARIOS = [
    ("timeout", "Claude request timed out"),
    ("rate_limit", "Rate limit exceeded"),
    ("auth_error", "Authentication failed"),
    ("network_error", "Network connection failed")
]