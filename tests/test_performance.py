"""
Performance tests for the Word Formatter system.
Tests system behavior under load and with large files.
"""

import os
import pytest
import time
import tempfile
from pathlib import Path
from unittest.mock import patch
from docx import Document

from document_converter import DocumentConverter
from document_converter_ai import AIDocumentConverter


class TestPerformanceBaseline:
    """Baseline performance tests for core functionality."""
    
    @pytest.mark.performance
    @pytest.mark.slow
    def test_style_extraction_performance(self, reference_doc_path):
        """Test style extraction performance with various document sizes."""
        from document_converter import StyleExtractor
        
        # Act
        start_time = time.time()
        extractor = StyleExtractor(str(reference_doc_path))
        end_time = time.time()
        
        # Assert
        extraction_time = end_time - start_time
        assert extraction_time < 5.0, f"Style extraction took too long: {extraction_time:.2f}s"
        
        # Verify styles were extracted
        assert extractor.styles is not None
        assert len(extractor.styles) > 0
    
    @pytest.mark.performance
    def test_small_file_conversion_performance(self, reference_doc_path, temp_dir):
        """Test performance with small files (baseline)."""
        # Arrange
        small_file = temp_dir / "small.txt"
        small_content = "Small Document\n\nChapter 1\n\nThis is a small test document with minimal content."
        small_file.write_text(small_content)
        
        converter = DocumentConverter(str(reference_doc_path))
        output_path = temp_dir / "small_output.docx"
        
        # Act
        start_time = time.time()
        converter.convert(str(small_file), str(output_path))
        end_time = time.time()
        
        # Assert
        conversion_time = end_time - start_time
        assert conversion_time < 2.0, f"Small file conversion took too long: {conversion_time:.2f}s"
        assert output_path.exists()
    
    @pytest.mark.performance
    @pytest.mark.slow
    def test_medium_file_conversion_performance(self, reference_doc_path, temp_dir):
        """Test performance with medium-sized files (~50KB)."""
        # Arrange
        medium_file = temp_dir / "medium.txt"
        content_block = "This is a paragraph with substantial content. " * 50
        chapters = []
        for i in range(20):
            chapters.append(f"Chapter {i+1}: Test Chapter\n\n{content_block}\n\n")
        
        medium_content = '\n'.join(chapters)
        medium_file.write_text(medium_content)
        
        converter = DocumentConverter(str(reference_doc_path))
        output_path = temp_dir / "medium_output.docx"
        
        # Act
        start_time = time.time()
        converter.convert(str(medium_file), str(output_path))
        end_time = time.time()
        
        # Assert
        conversion_time = end_time - start_time
        assert conversion_time < 10.0, f"Medium file conversion took too long: {conversion_time:.2f}s"
        assert output_path.exists()
        
        # Verify content integrity
        output_doc = Document(str(output_path))
        assert len(output_doc.paragraphs) >= 20


class TestLargeFilePerformance:
    """Performance tests for large file handling."""
    
    @pytest.mark.performance
    @pytest.mark.slow
    def test_large_file_conversion_performance(self, reference_doc_path, temp_dir):
        """Test performance with large files (~500KB)."""
        # Arrange
        large_file = temp_dir / "large.txt"
        content_block = "This is a substantial paragraph with meaningful content that will be repeated. " * 100
        chapters = []
        for i in range(50):
            chapters.append(f"Chapter {i+1}: Test Chapter\n\n{content_block}\n\n")
        
        large_content = '\n'.join(chapters)
        large_file.write_text(large_content)
        
        converter = DocumentConverter(str(reference_doc_path))
        output_path = temp_dir / "large_output.docx"
        
        # Act
        start_time = time.time()
        converter.convert(str(large_file), str(output_path))
        end_time = time.time()
        
        # Assert
        conversion_time = end_time - start_time
        assert conversion_time < 30.0, f"Large file conversion took too long: {conversion_time:.2f}s"
        assert output_path.exists()
        
        # Verify content integrity
        output_doc = Document(str(output_path))
        assert len(output_doc.paragraphs) >= 50
    
    @pytest.mark.performance
    @pytest.mark.slow
    @pytest.mark.claude
    def test_ai_large_file_chunking_performance(self, reference_doc_path, temp_dir, mock_claude_success, mock_notifications):
        """Test AI processing performance with chunking for large files."""
        # Arrange
        large_file = temp_dir / "ai_large.txt"
        content_block = "This is content that will require AI processing and chunking. " * 200
        chapters = []
        for i in range(20):
            chapters.append(f"Chapter {i+1}: AI Test Chapter\n\n{content_block}\n\n")
        
        large_content = '\n'.join(chapters)
        large_file.write_text(large_content)
        
        converter = AIDocumentConverter(str(reference_doc_path))
        converter.show_progress = False  # Disable progress for performance measurement
        output_path = temp_dir / "ai_large_output.docx"
        
        # Act
        start_time = time.time()
        result = converter.convert_with_ai(str(large_file), str(output_path))
        end_time = time.time()
        
        # Assert
        conversion_time = end_time - start_time
        assert conversion_time < 60.0, f"AI large file conversion took too long: {conversion_time:.2f}s"
        assert result is True
        assert output_path.exists()
    
    @pytest.mark.performance
    @pytest.mark.slow
    @pytest.mark.claude
    def test_chunking_threshold_performance(self, reference_doc_path, temp_dir, mock_claude_success, mock_notifications):
        """Test performance with different chunking thresholds."""
        # Arrange
        test_file = temp_dir / "chunk_test.txt"
        content = "Test content for chunking analysis. " * 1000  # ~40KB
        test_file.write_text(content)
        
        converter = AIDocumentConverter(str(reference_doc_path))
        converter.show_progress = False
        
        # Test with different chunk sizes
        chunk_sizes = [5000, 10000, 20000, 40000]
        results = {}
        
        for chunk_size in chunk_sizes:
            output_path = temp_dir / f"chunk_output_{chunk_size}.docx"
            
            # Act
            start_time = time.time()
            with patch.dict(os.environ, {'CHUNK_THRESHOLD': str(chunk_size)}):
                converter_test = AIDocumentConverter(str(reference_doc_path))
                converter_test.show_progress = False
                result = converter_test.convert_with_ai(str(test_file), str(output_path))
            end_time = time.time()
            
            # Store results
            results[chunk_size] = {
                'time': end_time - start_time,
                'success': result,
                'output_exists': output_path.exists()
            }
        
        # Assert
        for chunk_size, result in results.items():
            assert result['success'], f"Chunking with size {chunk_size} failed"
            assert result['output_exists'], f"Output file missing for chunk size {chunk_size}"
            assert result['time'] < 30.0, f"Chunking with size {chunk_size} took too long: {result['time']:.2f}s"


class TestMemoryPerformance:
    """Tests for memory usage and efficiency."""
    
    @pytest.mark.performance
    @pytest.mark.slow
    def test_memory_usage_large_files(self, reference_doc_path, temp_dir):
        """Test memory usage with large files."""
        import psutil
        import gc
        
        # Arrange
        large_file = temp_dir / "memory_test.txt"
        # Create a 1MB+ file
        content_block = "Memory usage test content. " * 1000
        chapters = []
        for i in range(100):
            chapters.append(f"Chapter {i+1}\n\n{content_block}\n\n")
        
        large_content = '\n'.join(chapters)
        large_file.write_text(large_content)
        
        # Measure initial memory
        process = psutil.Process()
        initial_memory = process.memory_info().rss / 1024 / 1024  # MB
        
        # Act
        converter = DocumentConverter(str(reference_doc_path))
        output_path = temp_dir / "memory_output.docx"
        converter.convert(str(large_file), str(output_path))
        
        # Measure peak memory
        peak_memory = process.memory_info().rss / 1024 / 1024  # MB
        memory_increase = peak_memory - initial_memory
        
        # Cleanup and measure final memory
        del converter
        gc.collect()
        final_memory = process.memory_info().rss / 1024 / 1024  # MB
        
        # Assert
        assert memory_increase < 100, f"Memory usage too high: {memory_increase:.2f}MB increase"
        assert output_path.exists()
        
        # Memory should be mostly released after cleanup
        memory_retained = final_memory - initial_memory
        assert memory_retained < 50, f"Too much memory retained: {memory_retained:.2f}MB"
    
    @pytest.mark.performance
    @pytest.mark.claude
    def test_ai_memory_usage_chunking(self, reference_doc_path, temp_dir, mock_claude_success, mock_notifications):
        """Test memory usage during AI chunking process."""
        import psutil
        import gc
        
        # Arrange
        chunking_file = temp_dir / "chunking_memory_test.txt"
        content = "Chunking memory test content. " * 2000  # Force chunking
        chunking_file.write_text(content)
        
        process = psutil.Process()
        initial_memory = process.memory_info().rss / 1024 / 1024  # MB
        
        # Act
        converter = AIDocumentConverter(str(reference_doc_path))
        converter.show_progress = False
        output_path = temp_dir / "chunking_memory_output.docx"
        
        result = converter.convert_with_ai(str(chunking_file), str(output_path))
        
        # Measure peak memory
        peak_memory = process.memory_info().rss / 1024 / 1024  # MB
        memory_increase = peak_memory - initial_memory
        
        # Cleanup
        del converter
        gc.collect()
        
        # Assert
        assert result is True
        assert memory_increase < 150, f"AI processing memory usage too high: {memory_increase:.2f}MB"


class TestConcurrentPerformance:
    """Tests for concurrent processing scenarios."""
    
    @pytest.mark.performance
    @pytest.mark.slow
    def test_multiple_converters_performance(self, reference_doc_path, temp_dir):
        """Test performance when multiple converters are used simultaneously."""
        import concurrent.futures
        
        # Arrange
        test_files = []
        for i in range(5):
            test_file = temp_dir / f"concurrent_test_{i}.txt"
            content = f"Concurrent test file {i}\n\n" + ("Test content. " * 100)
            test_file.write_text(content)
            test_files.append(test_file)
        
        def convert_file(file_path):
            converter = DocumentConverter(str(reference_doc_path))
            output_path = file_path.parent / f"{file_path.stem}_output.docx"
            start_time = time.time()
            converter.convert(str(file_path), str(output_path))
            end_time = time.time()
            return end_time - start_time, output_path.exists()
        
        # Act
        start_time = time.time()
        with concurrent.futures.ThreadPoolExecutor(max_workers=3) as executor:
            results = list(executor.map(convert_file, test_files))
        end_time = time.time()
        
        # Assert
        total_time = end_time - start_time
        assert total_time < 20.0, f"Concurrent processing took too long: {total_time:.2f}s"
        
        for conversion_time, file_exists in results:
            assert conversion_time < 10.0, f"Individual conversion took too long: {conversion_time:.2f}s"
            assert file_exists, "Output file was not created"
    
    @pytest.mark.performance 
    @pytest.mark.slow
    @pytest.mark.claude
    def test_concurrent_ai_processing_performance(self, reference_doc_path, temp_dir, mock_claude_success, mock_notifications):
        """Test performance with concurrent AI processing."""
        import concurrent.futures
        
        # Arrange
        test_files = []
        for i in range(3):  # Fewer files for AI processing
            test_file = temp_dir / f"ai_concurrent_test_{i}.txt"
            content = f"AI concurrent test file {i}\n\n" + ("AI test content. " * 200)
            test_file.write_text(content)
            test_files.append(test_file)
        
        def ai_convert_file(file_path):
            converter = AIDocumentConverter(str(reference_doc_path))
            converter.show_progress = False
            output_path = file_path.parent / f"{file_path.stem}_ai_output.docx"
            start_time = time.time()
            result = converter.convert_with_ai(str(file_path), str(output_path))
            end_time = time.time()
            return end_time - start_time, result, output_path.exists()
        
        # Act
        start_time = time.time()
        with concurrent.futures.ThreadPoolExecutor(max_workers=2) as executor:
            results = list(executor.map(ai_convert_file, test_files))
        end_time = time.time()
        
        # Assert
        total_time = end_time - start_time
        assert total_time < 30.0, f"Concurrent AI processing took too long: {total_time:.2f}s"
        
        for conversion_time, success, file_exists in results:
            assert success, "AI conversion should succeed"
            assert conversion_time < 15.0, f"Individual AI conversion took too long: {conversion_time:.2f}s"
            assert file_exists, "Output file was not created"


class TestScalabilityPerformance:
    """Tests for system scalability under various loads."""
    
    @pytest.mark.performance
    @pytest.mark.slow
    def test_progressive_file_sizes(self, reference_doc_path, temp_dir):
        """Test performance scaling with progressively larger files."""
        # File sizes in KB (approximate)
        file_sizes = [1, 5, 10, 25, 50, 100]
        results = {}
        
        for size_kb in file_sizes:
            # Arrange
            test_file = temp_dir / f"scale_test_{size_kb}kb.txt"
            # Calculate content to approximate the target size
            content_per_kb = "Test content for scalability analysis. " * 25  # ~1KB
            content = content_per_kb * size_kb
            test_file.write_text(content)
            
            converter = DocumentConverter(str(reference_doc_path))
            output_path = temp_dir / f"scale_output_{size_kb}kb.docx"
            
            # Act
            start_time = time.time()
            converter.convert(str(test_file), str(output_path))
            end_time = time.time()
            
            # Store results
            conversion_time = end_time - start_time
            results[size_kb] = {
                'time': conversion_time,
                'output_exists': output_path.exists(),
                'file_size_actual': test_file.stat().st_size / 1024  # Actual KB
            }
        
        # Assert
        for size_kb, result in results.items():
            assert result['output_exists'], f"Output missing for {size_kb}KB file"
            # Performance should scale reasonably (not exponentially)
            expected_max_time = min(30.0, size_kb * 0.5)  # Max 30s, or 0.5s per KB
            assert result['time'] < expected_max_time, \
                f"Conversion of {size_kb}KB file took too long: {result['time']:.2f}s"
        
        # Check that performance scales reasonably
        times = [results[size]['time'] for size in sorted(results.keys())]
        # Later files shouldn't take exponentially longer than earlier ones
        if len(times) >= 3:
            # Compare largest to smallest (shouldn't be more than 50x difference)
            ratio = times[-1] / times[0] if times[0] > 0 else float('inf')
            assert ratio < 50, f"Performance scaling issue: {ratio:.2f}x difference in processing time"