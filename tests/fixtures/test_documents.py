"""
Test document fixtures and sample data generators.
Creates various test documents for comprehensive testing.
"""

import tempfile
from pathlib import Path
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE


def create_comprehensive_reference_doc(output_path: Path) -> Path:
    """Create a comprehensive reference document with various styles."""
    doc = Document()
    
    # Set document margins
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1.25)
    section.right_margin = Inches(1.25)
    
    # Title
    title = doc.add_heading("Reference Document Template", level=0)
    title_run = title.runs[0]
    title_run.font.name = "Times New Roman"
    title_run.font.size = Pt(18)
    title_run.font.bold = True
    title_run.font.color.rgb = RGBColor(0, 0, 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Subtitle
    subtitle = doc.add_paragraph("A comprehensive formatting reference")
    subtitle_run = subtitle.runs[0]
    subtitle_run.font.name = "Times New Roman"
    subtitle_run.font.size = Pt(12)
    subtitle_run.font.italic = True
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.paragraph_format.space_after = Pt(18)
    
    # Chapter heading
    chapter = doc.add_heading("Chapter 1: Typography Standards", level=1)
    chapter_run = chapter.runs[0]
    chapter_run.font.name = "Times New Roman"
    chapter_run.font.size = Pt(16)
    chapter_run.font.bold = True
    chapter_run.font.color.rgb = RGBColor(0, 0, 139)  # Dark blue
    chapter.paragraph_format.space_before = Pt(24)
    chapter.paragraph_format.space_after = Pt(12)
    chapter.paragraph_format.keep_with_next = True
    
    # Section heading
    section_head = doc.add_heading("Section 1.1: Body Text Formatting", level=2)
    section_run = section_head.runs[0]
    section_run.font.name = "Times New Roman"
    section_run.font.size = Pt(14)
    section_run.font.bold = True
    section_run.font.color.rgb = RGBColor(0, 0, 0)
    section_head.paragraph_format.space_before = Pt(18)
    section_head.paragraph_format.space_after = Pt(6)
    
    # Normal paragraphs with various formatting
    para1 = doc.add_paragraph()
    para1.paragraph_format.space_after = Pt(12)
    para1.paragraph_format.line_spacing = 1.15
    para1.paragraph_format.first_line_indent = Pt(36)  # 0.5 inch
    
    run1 = para1.add_run("This is the standard body text format using ")
    run1.font.name = "Times New Roman"
    run1.font.size = Pt(11)
    
    run2 = para1.add_run("Times New Roman 11pt")
    run2.font.name = "Times New Roman"
    run2.font.size = Pt(11)
    run2.font.bold = True
    
    run3 = para1.add_run(" with proper spacing and indentation. This paragraph demonstrates the baseline formatting that should be applied to all body text.")
    run3.font.name = "Times New Roman"
    run3.font.size = Pt(11)
    
    # Paragraph with italic formatting
    para2 = doc.add_paragraph()
    para2.paragraph_format.space_after = Pt(12)
    para2.paragraph_format.line_spacing = 1.15
    para2.paragraph_format.first_line_indent = Pt(36)
    
    run4 = para2.add_run("Foreign language terms should be ")
    run4.font.name = "Times New Roman"
    run4.font.size = Pt(11)
    
    run5 = para2.add_run("italicized")
    run5.font.name = "Times New Roman"
    run5.font.size = Pt(11)
    run5.font.italic = True
    
    run6 = para2.add_run(", such as ")
    run6.font.name = "Times New Roman"
    run6.font.size = Pt(11)
    
    run7 = para2.add_run("mokṣa")
    run7.font.name = "Times New Roman"
    run7.font.size = Pt(11)
    run7.font.italic = True
    
    run8 = para2.add_run(" and ")
    run8.font.name = "Times New Roman"
    run8.font.size = Pt(11)
    
    run9 = para2.add_run("dharma")
    run9.font.name = "Times New Roman"
    run9.font.size = Pt(11)
    run9.font.italic = True
    
    run10 = para2.add_run(" from Sanskrit.")
    run10.font.name = "Times New Roman"
    run10.font.size = Pt(11)
    
    # Subsection
    subsection = doc.add_heading("Subsection 1.1.1: Special Formatting", level=3)
    subsection_run = subsection.runs[0]
    subsection_run.font.name = "Times New Roman"
    subsection_run.font.size = Pt(12)
    subsection_run.font.bold = True
    subsection.paragraph_format.space_before = Pt(12)
    subsection.paragraph_format.space_after = Pt(6)
    
    # Quote paragraph
    quote = doc.add_paragraph()
    quote.paragraph_format.left_indent = Pt(72)  # 1 inch
    quote.paragraph_format.right_indent = Pt(36)  # 0.5 inch
    quote.paragraph_format.space_before = Pt(12)
    quote.paragraph_format.space_after = Pt(12)
    quote.paragraph_format.line_spacing = 1.0
    
    quote_run = quote.add_run("This is a block quote format used for longer quotations. It should be indented from both margins and use single spacing.")
    quote_run.font.name = "Times New Roman"
    quote_run.font.size = Pt(10)
    
    # Bullet list
    list_intro = doc.add_paragraph("Key formatting principles include:")
    list_intro.paragraph_format.space_after = Pt(6)
    list_intro_run = list_intro.runs[0]
    list_intro_run.font.name = "Times New Roman"
    list_intro_run.font.size = Pt(11)
    
    bullet_items = [
        "Consistent font usage throughout the document",
        "Proper heading hierarchy and spacing",
        "Appropriate line spacing and paragraph breaks",
        "Correct indentation for body paragraphs"
    ]
    
    for item in bullet_items:
        bullet_para = doc.add_paragraph(item, style='List Bullet')
        bullet_run = bullet_para.runs[0]
        bullet_run.font.name = "Times New Roman"
        bullet_run.font.size = Pt(11)
    
    # Numbered list
    doc.add_paragraph("Document structure should follow this order:", style='Normal')
    
    numbered_items = [
        "Title page with proper formatting",
        "Table of contents (if applicable)",
        "Main body text with consistent styling", 
        "References and bibliography"
    ]
    
    for i, item in enumerate(numbered_items, 1):
        num_para = doc.add_paragraph(f"{i}. {item}", style='List Number')
        num_run = num_para.runs[0]
        num_run.font.name = "Times New Roman"
        num_run.font.size = Pt(11)
    
    # Table
    doc.add_paragraph("Sample formatting table:", style='Normal')
    
    table = doc.add_table(rows=4, cols=3)
    table.style = 'Table Grid'
    
    # Header row
    header_cells = table.rows[0].cells
    header_cells[0].text = "Element"
    header_cells[1].text = "Font"
    header_cells[2].text = "Size"
    
    # Make header bold
    for cell in header_cells:
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.font.bold = True
                run.font.name = "Times New Roman"
                run.font.size = Pt(11)
    
    # Data rows
    data_rows = [
        ["Title", "Times New Roman", "18pt"],
        ["Heading 1", "Times New Roman", "16pt"],
        ["Body Text", "Times New Roman", "11pt"]
    ]
    
    for i, row_data in enumerate(data_rows, 1):
        row_cells = table.rows[i].cells
        for j, cell_text in enumerate(row_data):
            row_cells[j].text = cell_text
            for paragraph in row_cells[j].paragraphs:
                for run in paragraph.runs:
                    run.font.name = "Times New Roman"
                    run.font.size = Pt(11)
    
    # Page break for new chapter
    doc.add_page_break()
    
    # Second chapter
    chapter2 = doc.add_heading("Chapter 2: Advanced Formatting", level=1)
    chapter2_run = chapter2.runs[0]
    chapter2_run.font.name = "Times New Roman"
    chapter2_run.font.size = Pt(16)
    chapter2_run.font.bold = True
    chapter2_run.font.color.rgb = RGBColor(0, 0, 139)
    chapter2.paragraph_format.space_before = Pt(24)
    chapter2.paragraph_format.space_after = Pt(12)
    
    # Final paragraph
    final_para = doc.add_paragraph("This reference document provides the baseline formatting that should be applied to all converted documents. The styles defined here ensure consistency and professional appearance.")
    final_run = final_para.runs[0]
    final_run.font.name = "Times New Roman"
    final_run.font.size = Pt(11)
    final_para.paragraph_format.space_after = Pt(12)
    final_para.paragraph_format.line_spacing = 1.15
    final_para.paragraph_format.first_line_indent = Pt(36)
    
    # Save document
    doc.save(str(output_path))
    return output_path


def create_sample_test_documents(base_dir: Path) -> dict:
    """Create a comprehensive set of sample documents for testing."""
    base_dir.mkdir(parents=True, exist_ok=True)
    
    documents = {}
    
    # 1. Simple text document
    simple_txt = base_dir / "simple_document.txt"
    simple_content = """Simple Test Document

Introduction

This is a basic text document for testing the conversion process. It contains simple paragraphs and basic structure.

Chapter 1: Getting Started

This chapter explains the basic concepts. The text should be converted to properly formatted Word document with consistent styling.

The converter should handle line breaks and paragraph spacing correctly.

Chapter 2: Advanced Topics  

This section covers more complex topics and demonstrates how the system handles multiple chapters.

Conclusion

This concludes our simple test document."""
    simple_txt.write_text(simple_content)
    documents['simple_text'] = simple_txt
    
    # 2. Complex text with various elements
    complex_txt = base_dir / "complex_document.txt"
    complex_content = """Complex Test Document

INTRODUCTION TO ADVANCED CONCEPTS

This document contains various text elements that need to be processed correctly.

Chapter 1: Lists and Structure

Here are some important points:

- First bullet point
- Second bullet point with more details
- Third point about formatting

Numbered list:
1. First numbered item
2. Second item with explanation
3. Final numbered point

Chapter 2: Special Characters and Formatting

This section contains text with special characters: "quotes", 'apostrophes', and—em dashes.

Mathematical symbols: α β γ δ ∑ ∏ ∫ ∞

Currency symbols: $ € £ ¥ ₹

Chapter 3: International Content

Sanskrit transliterations: Vāg vai brahma | vācā hy evedaṃ sarvaṃ sṛṣṭam

Devanagari script: हिन्दी परीक्षण पाठ

Arabic text: السلام عليكم

Chinese characters: 你好世界

Chapter 4: Technical Content

Code-like text:
function example() {
    return "formatted text";
}

File paths: /usr/local/bin/application

URLs: https://www.example.com

Conclusion

This complex document tests various aspects of text processing and conversion."""
    complex_txt.write_text(complex_content)
    documents['complex_text'] = complex_txt
    
    # 3. Markdown document
    markdown_md = base_dir / "sample_markdown.md"
    markdown_content = """# Markdown Test Document

## Introduction

This is a **Markdown** document with various formatting elements.

### Lists

Unordered list:
- Item one
- Item two  
- Item three

Ordered list:
1. First item
2. Second item
3. Third item

### Text Formatting

This paragraph contains *italic text*, **bold text**, and `inline code`.

### Blockquotes

> This is a blockquote with important information.
> It can span multiple lines.

### Code Blocks

```python
def hello_world():
    print("Hello, World!")
    return True
```

### Tables

| Column 1 | Column 2 | Column 3 |
|----------|----------|----------|
| Data A   | Data B   | Data C   |
| Data D   | Data E   | Data F   |

### Links and Images

Visit [OpenAI](https://openai.com) for more information.

### Horizontal Rule

---

## Conclusion

This Markdown document demonstrates various formatting elements that should be converted properly to Word format."""
    markdown_md.write_text(markdown_content)
    documents['markdown'] = markdown_md
    
    # 4. RTF document
    rtf_file = base_dir / "sample_rtf.rtf"
    rtf_content = r"""{\rtf1\ansi\deff0 {\fonttbl {\f0 Times New Roman;}}
{\colortbl ;\red0\green0\blue0;}
\f0\fs22\cf1 
RTF Test Document\par
\par
This is a Rich Text Format (RTF) document for testing.\par
\par
\b Bold text\b0  and \i italic text\i0  should be preserved.\par
\par
Multiple paragraphs should be handled correctly.\par
\par
Special characters: \u233\'e9 \u224\'e0 \u231\'e7\par
\par
End of RTF test document.\par
}"""
    rtf_file.write_text(rtf_content)
    documents['rtf'] = rtf_file
    
    # 5. Large document for performance testing
    large_txt = base_dir / "large_document.txt"
    large_content_parts = []
    
    large_content_parts.append("Large Document Performance Test\n\n")
    
    for chapter in range(1, 26):  # 25 chapters
        large_content_parts.append(f"Chapter {chapter}: Performance Testing\n\n")
        
        # Add substantial content to each chapter
        for section in range(1, 6):  # 5 sections per chapter
            large_content_parts.append(f"Section {chapter}.{section}: Content Analysis\n\n")
            
            # Multiple paragraphs per section
            for para in range(1, 6):  # 5 paragraphs per section
                paragraph = f"This is paragraph {para} of section {section} in chapter {chapter}. " + \
                           "It contains substantial text content to test performance with large documents. " * 10
                large_content_parts.append(paragraph + "\n\n")
    
    large_content_parts.append("Conclusion\n\nThis large document tests the system's ability to handle substantial content efficiently.")
    
    large_content = ''.join(large_content_parts)
    large_txt.write_text(large_content)
    documents['large_text'] = large_txt
    
    # 6. Unicode and international text
    unicode_txt = base_dir / "unicode_document.txt"
    unicode_content = """Unicode and International Text Test

English Section
This section contains regular English text.

European Languages
French: Bonjour, comment allez-vous? Café, naïve, résumé.
German: Guten Tag, wie geht es Ihnen? Müller, Größe, Weiß.
Spanish: Hola, ¿cómo está usted? Niño, señor, año.
Italian: Ciao, come stai? Città, più, perché.

Slavic Languages  
Russian: Привет, как дела? Москва, Россия, спасибо.
Polish: Cześć, jak się masz? Warszawa, Polska, dziękuję.

Asian Languages
Chinese (Simplified): 你好，你好吗？北京，中国，谢谢。
Chinese (Traditional): 你好，你好嗎？台北，台灣，謝謝。
Japanese: こんにちは、元気ですか？東京、日本、ありがとう。
Korean: 안녕하세요, 잘 지내세요? 서울, 한국, 감사합니다.

Indian Languages
Hindi: नमस्ते, आप कैसे हैं? दिल्ली, भारत, धन्यवाद।
Sanskrit: नमस्ते। सत्यमेव जयते। धर्मो रक्षति रक्षितः।
Sanskrit Transliteration: Namasté. Satyameva jayate. Dharmo rakṣati rakṣitaḥ.

Middle Eastern Languages
Arabic: السلام عليكم، كيف حالك؟ القاهرة، مصر، شكراً.
Hebrew: שלום, מה שלומך? ירושלים, ישראל, תודה.
Persian: سلام، حال شما چطور است؟ تهران، ایران، متشکرم.

Southeast Asian Languages
Thai: สวัสดี คุณสบายดีไหม? กรุงเทพ ประเทศไทย ขอบคุณ
Vietnamese: Xin chào, bạn có khỏe không? Hà Nội, Việt Nam, cảm ơn.

African Languages
Swahili: Hujambo, habari gani? Nairobi, Kenya, asante.

Mathematical and Technical Symbols
Greek letters: α β γ δ ε ζ η θ ι κ λ μ ν ξ ο π ρ σ τ υ φ χ ψ ω
Mathematical: ∑ ∏ ∫ ∞ ≈ ≠ ≤ ≥ ± ∓ × ÷ √ ∂ ∇
Currency: $ € £ ¥ ₹ ₽ ₩ ₫ ₪ ₦ ₨
Arrows: ← → ↑ ↓ ↔ ⇒ ⇔ ⇐ ⇑ ⇓

Emoji and Symbols
Faces: 😀 😃 😄 😁 😆 😅 😂 🤣
Objects: 🚀 💻 📱 📄 📝 🔧 ⚙️ 🌟
Nature: 🌍 🌎 🌏 🌞 🌛 ⭐ 🌈 🌸

Mixed Content Examples
The Sanskrit concept of mokṣa (मोक्ष) is central to Hindu philosophy.
In Arabic, peace is سلام (salām), similar to Hebrew שלום (shalom).
The Chinese character 道 (dào) represents "the way" or "path".
Mathematical expressions: E = mc² and ∫f(x)dx represent fundamental concepts.

This document tests comprehensive Unicode support across multiple writing systems and symbol sets."""
    unicode_txt.write_text(unicode_content, encoding='utf-8')
    documents['unicode'] = unicode_txt
    
    # 7. Empty and minimal documents
    empty_txt = base_dir / "empty_document.txt"
    empty_txt.write_text("")
    documents['empty'] = empty_txt
    
    minimal_txt = base_dir / "minimal_document.txt"
    minimal_txt.write_text("Minimal Document\n\nJust one paragraph.")
    documents['minimal'] = minimal_txt
    
    # 8. Malformed content
    malformed_txt = base_dir / "malformed_document.txt"
    malformed_content = """Malformed Test Document


Multiple empty lines above.

    Inconsistent indentation.
\tTabs and spaces mixed.

Lines with trailing spaces:    
More trailing spaces:        

SHOUTING TEXT EVERYWHERE

mixed Case Words Throughout

"Unmatched quotes and 'apostrophes

Line breaks in weird
places that might
confuse processing

!!!Multiple!!! punctuation??? marks...

(Parentheses without closing
[Brackets without closing
{Braces without closing

End of malformed content."""
    malformed_txt.write_text(malformed_content)
    documents['malformed'] = malformed_txt
    
    return documents


def create_test_reference_formats(base_dir: Path) -> dict:
    """Create various reference format documents for testing different styles."""
    base_dir.mkdir(parents=True, exist_ok=True)
    
    references = {}
    
    # 1. Standard academic format
    academic_ref = base_dir / "academic_reference.docx"
    academic_ref = create_comprehensive_reference_doc(academic_ref)
    references['academic'] = academic_ref
    
    # 2. Simple format with minimal styling
    simple_ref = base_dir / "simple_reference.docx"
    doc = Document()
    
    # Set basic margins
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    
    # Simple title
    title = doc.add_heading("Simple Reference", level=0)
    title.runs[0].font.name = "Arial"
    title.runs[0].font.size = Pt(14)
    
    # Simple paragraph
    para = doc.add_paragraph("This is a simple reference format with minimal styling.")
    para.runs[0].font.name = "Arial"
    para.runs[0].font.size = Pt(10)
    
    doc.save(str(simple_ref))
    references['simple'] = simple_ref
    
    # 3. Complex format with detailed styling
    complex_ref = base_dir / "complex_reference.docx"
    doc = Document()
    
    # Custom margins
    section = doc.sections[0]
    section.top_margin = Inches(1.5)
    section.bottom_margin = Inches(1.25)
    section.left_margin = Inches(1.75)
    section.right_margin = Inches(1.25)
    
    # Title with custom formatting
    title = doc.add_heading("Complex Reference Format", level=0)
    title.runs[0].font.name = "Garamond"
    title.runs[0].font.size = Pt(20)
    title.runs[0].font.color.rgb = RGBColor(139, 0, 0)  # Dark red
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Subtitle
    subtitle = doc.add_paragraph("Advanced Formatting Specifications")
    subtitle.runs[0].font.name = "Garamond"
    subtitle.runs[0].font.size = Pt(14)
    subtitle.runs[0].font.italic = True
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Heading with detailed formatting
    heading = doc.add_heading("Chapter 1: Specifications", level=1)
    heading.runs[0].font.name = "Garamond"
    heading.runs[0].font.size = Pt(16)
    heading.runs[0].font.bold = True
    heading.runs[0].font.color.rgb = RGBColor(0, 100, 0)  # Dark green
    heading.paragraph_format.space_before = Pt(30)
    heading.paragraph_format.space_after = Pt(15)
    
    # Body text with custom formatting
    para = doc.add_paragraph()
    para.paragraph_format.line_spacing = 1.25
    para.paragraph_format.space_after = Pt(15)
    para.paragraph_format.first_line_indent = Pt(50)
    
    run = para.add_run("This reference document uses custom Garamond font with specific spacing and color scheme for professional document appearance.")
    run.font.name = "Garamond"
    run.font.size = Pt(12)
    
    doc.save(str(complex_ref))
    references['complex'] = complex_ref
    
    return references


if __name__ == "__main__":
    # Create test fixtures when run directly
    fixtures_dir = Path(__file__).parent
    
    # Create sample documents
    docs_dir = fixtures_dir / "sample_docs"
    sample_docs = create_sample_test_documents(docs_dir)
    print(f"Created {len(sample_docs)} sample documents in {docs_dir}")
    
    # Create reference formats
    ref_dir = fixtures_dir / "reference_formats"
    reference_docs = create_test_reference_formats(ref_dir)
    print(f"Created {len(reference_docs)} reference documents in {ref_dir}")
    
    print("\nTest fixtures created successfully!")
    print("Sample documents:")
    for name, path in sample_docs.items():
        print(f"  {name}: {path}")
    
    print("\nReference documents:")
    for name, path in reference_docs.items():
        print(f"  {name}: {path}")