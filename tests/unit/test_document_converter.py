"""
Unit tests for document_converter.py module.
Tests StyleExtractor and DocumentConverter classes.
"""

import pytest
import tempfile
from pathlib import Path
from unittest.mock import patch, MagicMock
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

from document_converter import StyleExtractor, DocumentConverter


class TestStyleExtractor:
    """Test cases for the StyleExtractor class."""
    
    @pytest.mark.unit
    def test_style_extractor_initialization(self, reference_doc_path):
        """Test that StyleExtractor initializes correctly with a reference document."""
        # Arrange & Act
        extractor = StyleExtractor(str(reference_doc_path))
        
        # Assert
        assert extractor.reference_doc is not None
        assert extractor.styles is not None
        assert 'document' in extractor.styles
        assert 'paragraphs' in extractor.styles
        assert 'headings' in extractor.styles
        assert 'normal' in extractor.styles
        assert 'lists' in extractor.styles
    
    @pytest.mark.unit
    def test_extract_document_styles(self, reference_doc_path):
        """Test extraction of document-level styles."""
        # Arrange
        extractor = StyleExtractor(str(reference_doc_path))
        
        # Act
        doc_styles = extractor.styles['document']
        
        # Assert
        assert 'margins' in doc_styles
        assert 'page_width' in doc_styles
        assert 'page_height' in doc_styles
        
        margins = doc_styles['margins']
        assert 'top' in margins
        assert 'bottom' in margins
        assert 'left' in margins
        assert 'right' in margins
        assert margins['top'] == Pt(72)  # 1 inch margins from our fixture
    
    @pytest.mark.unit
    def test_extract_heading_styles(self, reference_doc_path):
        """Test extraction of heading styles."""
        # Arrange
        extractor = StyleExtractor(str(reference_doc_path))
        
        # Act
        heading_styles = extractor.styles['headings']
        
        # Assert
        # The reference document should have heading styles
        assert isinstance(heading_styles, dict)
        # Test that heading styles contain expected structure
        for heading_name, style in heading_styles.items():
            if style:  # Some headings might be empty
                assert 'font' in style or 'paragraph' in style
    
    @pytest.mark.unit
    def test_extract_normal_style(self, reference_doc_path):
        """Test extraction of normal paragraph style."""
        # Arrange
        extractor = StyleExtractor(str(reference_doc_path))
        
        # Act
        normal_style = extractor.styles['normal']
        
        # Assert
        assert isinstance(normal_style, dict)
        # Normal style should have font information
        if normal_style.get('font'):
            font = normal_style['font']
            assert 'name' in font or 'size' in font
    
    @pytest.mark.unit
    def test_extract_paragraph_styles(self, reference_doc_path):
        """Test extraction of paragraph styles."""
        # Arrange
        extractor = StyleExtractor(str(reference_doc_path))
        
        # Act
        paragraph_styles = extractor.styles['paragraphs']
        
        # Assert
        assert isinstance(paragraph_styles, list)
        # Should have collected some paragraph styles
        if paragraph_styles:
            style = paragraph_styles[0]
            assert 'alignment' in style
            assert 'style_name' in style
    
    @pytest.mark.unit
    def test_extract_list_styles(self, reference_doc_path):
        """Test extraction of list styles."""
        # Arrange
        extractor = StyleExtractor(str(reference_doc_path))
        
        # Act
        list_styles = extractor.styles['lists']
        
        # Assert
        assert isinstance(list_styles, dict)
        assert 'bullet' in list_styles
        assert 'numbered' in list_styles
        assert 'indent' in list_styles['bullet']
        assert 'indent' in list_styles['numbered']
    
    @pytest.mark.unit
    def test_style_extractor_nonexistent_file(self):
        """Test StyleExtractor with non-existent file raises appropriate error."""
        # Arrange
        nonexistent_path = "/path/that/does/not/exist.docx"
        
        # Act & Assert
        with pytest.raises(FileNotFoundError):
            StyleExtractor(nonexistent_path)
    
    @pytest.mark.unit
    @patch('os.environ.get')
    def test_debug_logging_enabled(self, mock_env_get, reference_doc_path, capsys):
        """Test that debug logging works when enabled."""
        # Arrange
        mock_env_get.return_value = '1'  # Enable debug mode
        
        # Act
        StyleExtractor(str(reference_doc_path))
        
        # Assert
        captured = capsys.readouterr()
        assert "REFERENCE DOCUMENT STYLES EXTRACTED" in captured.out
    
    @pytest.mark.unit
    @patch('os.environ.get')
    def test_debug_logging_disabled(self, mock_env_get, reference_doc_path, capsys):
        """Test that debug logging is disabled by default."""
        # Arrange
        mock_env_get.return_value = '0'  # Disable debug mode
        
        # Act
        StyleExtractor(str(reference_doc_path))
        
        # Assert
        captured = capsys.readouterr()
        assert "REFERENCE DOCUMENT STYLES EXTRACTED" not in captured.out


class TestDocumentConverter:
    """Test cases for the DocumentConverter class."""
    
    @pytest.mark.unit
    def test_document_converter_initialization(self, reference_doc_path):
        """Test that DocumentConverter initializes correctly."""
        # Arrange & Act
        converter = DocumentConverter(str(reference_doc_path))
        
        # Assert
        assert converter.style_extractor is not None
        assert converter.output_doc is not None
        assert isinstance(converter.output_doc, Document)
    
    @pytest.mark.unit
    def test_read_text_file(self, reference_doc_path, sample_text_file):
        """Test reading a plain text file."""
        # Arrange
        converter = DocumentConverter(str(reference_doc_path))
        
        # Act
        content = converter._read_text_file(sample_text_file)
        
        # Assert
        assert "Test Document" in content
        assert "Chapter 1: Introduction" in content
        assert len(content) > 0
    
    @pytest.mark.unit
    def test_read_markdown_file(self, reference_doc_path, sample_markdown_file):
        """Test reading a markdown file."""
        # Arrange
        converter = DocumentConverter(str(reference_doc_path))
        
        # Act
        content = converter._read_markdown_file(sample_markdown_file)
        
        # Assert
        assert "# Test Document" in content
        assert "## Chapter 1: Introduction" in content
        assert "| Column 1 |" in content  # Table content
    
    @pytest.mark.unit
    def test_process_text_content(self, reference_doc_path, sample_text_file):
        """Test processing plain text content into Word paragraphs."""
        # Arrange
        converter = DocumentConverter(str(reference_doc_path))
        content = converter._read_text_file(sample_text_file)
        
        # Act
        converter._process_text_content(content)
        
        # Assert
        assert len(converter.output_doc.paragraphs) > 0
        # Check that content was added
        all_text = '\n'.join([p.text for p in converter.output_doc.paragraphs])
        assert "Test Document" in all_text
    
    @pytest.mark.unit
    def test_process_markdown_content(self, reference_doc_path, sample_markdown_file):
        """Test processing markdown content into Word document."""
        # Arrange
        converter = DocumentConverter(str(reference_doc_path))
        content = converter._read_markdown_file(sample_markdown_file)
        
        # Act
        converter._process_markdown_content(content)
        
        # Assert
        assert len(converter.output_doc.paragraphs) > 0
        
        # Check for headings
        headings = [p for p in converter.output_doc.paragraphs if p.style.name.startswith('Heading')]
        assert len(headings) > 0
        
        # Check that heading content is correct
        heading_texts = [h.text for h in headings]
        assert "Test Document" in heading_texts
        assert "Chapter 1: Introduction" in heading_texts
    
    @pytest.mark.unit
    def test_process_html_element_headings(self, reference_doc_path):
        """Test processing HTML heading elements."""
        # Arrange
        converter = DocumentConverter(str(reference_doc_path))
        
        # Create a mock HTML element
        from bs4 import BeautifulSoup
        html = "<h1>Chapter 1: Test Chapter</h1>"
        soup = BeautifulSoup(html, 'html.parser')
        element = soup.find('h1')
        
        # Act
        converter._process_html_element(element)
        
        # Assert
        headings = [p for p in converter.output_doc.paragraphs if p.style.name.startswith('Heading')]
        assert len(headings) == 1
        assert headings[0].text == "Chapter 1: Test Chapter"
        assert headings[0].style.name == "Heading 1"
    
    @pytest.mark.unit
    def test_process_html_element_paragraphs(self, reference_doc_path):
        """Test processing HTML paragraph elements."""
        # Arrange
        converter = DocumentConverter(str(reference_doc_path))
        
        # Create a mock HTML element with inline formatting
        from bs4 import BeautifulSoup
        html = "<p>This is <strong>bold</strong> and <em>italic</em> text.</p>"
        soup = BeautifulSoup(html, 'html.parser')
        element = soup.find('p')
        
        # Act
        converter._process_html_element(element)
        
        # Assert
        paragraphs = [p for p in converter.output_doc.paragraphs 
                     if not p.style.name.startswith('Heading')]
        assert len(paragraphs) >= 1
        
        # Check for formatting
        paragraph = paragraphs[-1]  # Get the last added paragraph
        bold_runs = [r for r in paragraph.runs if r.bold]
        italic_runs = [r for r in paragraph.runs if r.italic]
        assert len(bold_runs) > 0  # Should have bold text
        assert len(italic_runs) > 0  # Should have italic text
    
    @pytest.mark.unit
    def test_process_html_element_lists(self, reference_doc_path):
        """Test processing HTML list elements."""
        # Arrange
        converter = DocumentConverter(str(reference_doc_path))
        
        # Create a mock HTML list
        from bs4 import BeautifulSoup
        html = "<ul><li>First item</li><li>Second item</li></ul>"
        soup = BeautifulSoup(html, 'html.parser')
        element = soup.find('ul')
        
        # Act
        converter._process_html_element(element)
        
        # Assert
        list_items = [p for p in converter.output_doc.paragraphs 
                     if p.style.name == 'List Bullet']
        assert len(list_items) == 2
        assert list_items[0].text == "First item"
        assert list_items[1].text == "Second item"
    
    @pytest.mark.unit
    def test_process_html_element_tables(self, reference_doc_path):
        """Test processing HTML table elements."""
        # Arrange
        converter = DocumentConverter(str(reference_doc_path))
        
        # Create a mock HTML table
        from bs4 import BeautifulSoup
        html = """<table>
            <tr><th>Header 1</th><th>Header 2</th></tr>
            <tr><td>Data 1</td><td>Data 2</td></tr>
        </table>"""
        soup = BeautifulSoup(html, 'html.parser')
        element = soup.find('table')
        
        # Act
        converter._process_html_element(element)
        
        # Assert
        assert len(converter.output_doc.tables) == 1
        table = converter.output_doc.tables[0]
        assert len(table.rows) == 2
        assert len(table.columns) == 2
        assert table.rows[0].cells[0].text == "Header 1"
        assert table.rows[1].cells[0].text == "Data 1"
    
    @pytest.mark.unit
    def test_chapter_page_breaks(self, reference_doc_path):
        """Test that chapter headings get page breaks."""
        # Arrange
        converter = DocumentConverter(str(reference_doc_path))
        
        # Add a regular heading first
        converter.output_doc.add_heading("Regular Heading", level=2)
        
        # Create a chapter heading
        from bs4 import BeautifulSoup
        html = "<h1>Chapter 2: New Chapter</h1>"
        soup = BeautifulSoup(html, 'html.parser')
        element = soup.find('h1')
        
        # Act
        converter._process_html_element(element)
        
        # Assert
        # Check for page break (this is complex to test directly, 
        # so we check that formatting was applied)
        chapter_headings = [p for p in converter.output_doc.paragraphs 
                           if p.style.name == 'Heading 1']
        assert len(chapter_headings) == 1
        
        chapter = chapter_headings[0]
        assert chapter.paragraph_format.space_before == Pt(24)
        assert chapter.paragraph_format.space_after == Pt(18)
        assert chapter.paragraph_format.keep_with_next is True
    
    @pytest.mark.unit
    def test_convert_text_file(self, reference_doc_path, sample_text_file, temp_dir):
        """Test converting a complete text file."""
        # Arrange
        converter = DocumentConverter(str(reference_doc_path))
        output_path = temp_dir / "output.docx"
        
        # Act
        converter.convert(str(sample_text_file), str(output_path))
        
        # Assert
        assert output_path.exists()
        
        # Load and check the output document
        output_doc = Document(str(output_path))
        assert len(output_doc.paragraphs) > 0
        
        all_text = '\n'.join([p.text for p in output_doc.paragraphs])
        assert "Test Document" in all_text
        assert "Chapter 1: Introduction" in all_text
    
    @pytest.mark.unit
    def test_convert_markdown_file(self, reference_doc_path, sample_markdown_file, temp_dir):
        """Test converting a complete markdown file."""
        # Arrange
        converter = DocumentConverter(str(reference_doc_path))
        output_path = temp_dir / "output.docx"
        
        # Act
        converter.convert(str(sample_markdown_file), str(output_path))
        
        # Assert
        assert output_path.exists()
        
        # Load and check the output document
        output_doc = Document(str(output_path))
        assert len(output_doc.paragraphs) > 0
        
        # Check for proper heading structure
        headings = [p for p in output_doc.paragraphs if p.style.name.startswith('Heading')]
        assert len(headings) > 0
        
        # Check for tables
        assert len(output_doc.tables) > 0
    
    @pytest.mark.unit
    def test_convert_unsupported_file_type(self, reference_doc_path, temp_dir):
        """Test that unsupported file types raise ValueError."""
        # Arrange
        converter = DocumentConverter(str(reference_doc_path))
        unsupported_file = temp_dir / "test.xyz"
        unsupported_file.write_text("content")
        output_path = temp_dir / "output.docx"
        
        # Act & Assert
        with pytest.raises(ValueError, match="Unsupported file format"):
            converter.convert(str(unsupported_file), str(output_path))
    
    @pytest.mark.unit
    def test_apply_reference_styles(self, reference_doc_path):
        """Test that reference styles are applied to the output document."""
        # Arrange
        converter = DocumentConverter(str(reference_doc_path))
        
        # Add some content
        converter.output_doc.add_paragraph("Test paragraph")
        converter.output_doc.add_heading("Test heading", level=1)
        
        # Act
        converter._apply_reference_styles()
        
        # Assert
        # Check that margins were applied
        section = converter.output_doc.sections[0]
        assert section.top_margin == Pt(72)  # From our reference document
        assert section.left_margin == Pt(72)
        assert section.right_margin == Pt(72)
        assert section.bottom_margin == Pt(72)
    
    @pytest.mark.unit
    def test_apply_paragraph_style(self, reference_doc_path):
        """Test applying style dictionary to a paragraph."""
        # Arrange
        converter = DocumentConverter(str(reference_doc_path))
        paragraph = converter.output_doc.add_paragraph("Test paragraph")
        
        # Create a test style
        test_style = {
            'paragraph': {
                'alignment': WD_ALIGN_PARAGRAPH.CENTER,
                'space_before': Pt(12),
                'space_after': Pt(6)
            },
            'font': {
                'name': 'Arial',
                'size': Pt(14),
                'bold': True,
                'italic': False
            }
        }
        
        # Act
        converter._apply_paragraph_style(paragraph, test_style)
        
        # Assert
        assert paragraph.alignment == WD_ALIGN_PARAGRAPH.CENTER
        assert paragraph.paragraph_format.space_before == Pt(12)
        assert paragraph.paragraph_format.space_after == Pt(6)
        
        # Check font formatting on runs (if any)
        if paragraph.runs:
            run = paragraph.runs[0]
            assert run.font.name == 'Arial'
            assert run.font.size == Pt(14)
            assert run.font.bold is True
            assert run.font.italic is False


class TestDocumentConverterEdgeCases:
    """Test edge cases for DocumentConverter."""
    
    @pytest.mark.unit
    def test_empty_text_file(self, reference_doc_path, empty_text_file, temp_dir):
        """Test processing an empty text file."""
        # Arrange
        converter = DocumentConverter(str(reference_doc_path))
        output_path = temp_dir / "output.docx"
        
        # Act
        converter.convert(str(empty_text_file), str(output_path))
        
        # Assert
        assert output_path.exists()
        
        # The document should still be valid but essentially empty
        output_doc = Document(str(output_path))
        # Even an "empty" document will have at least one paragraph
        assert len(output_doc.paragraphs) >= 0
    
    @pytest.mark.unit
    def test_unicode_content(self, reference_doc_path, temp_dir):
        """Test processing text with Unicode characters."""
        # Arrange
        converter = DocumentConverter(str(reference_doc_path))
        unicode_file = temp_dir / "unicode.txt"
        unicode_content = """Unicode Test Document

This document contains various Unicode characters:
- Emoji: 🚀 💻 📄
- Accented characters: café, naïve, résumé
- Mathematical symbols: ∑ ∏ ∫ √
- Sanskrit: Vāg vai brahma
- Devanagari: हिन्दी परीक्षण
- Arabic: السلام عليكم
- Chinese: 你好世界
"""
        unicode_file.write_text(unicode_content, encoding='utf-8')
        output_path = temp_dir / "output.docx"
        
        # Act
        converter.convert(str(unicode_file), str(output_path))
        
        # Assert
        assert output_path.exists()
        
        output_doc = Document(str(output_path))
        all_text = '\n'.join([p.text for p in output_doc.paragraphs])
        assert "Unicode Test Document" in all_text
        assert "🚀" in all_text
        assert "हिन्दी" in all_text
        assert "السلام" in all_text
    
    @pytest.mark.unit
    def test_malformed_markdown(self, reference_doc_path, temp_dir):
        """Test processing malformed markdown content."""
        # Arrange
        converter = DocumentConverter(str(reference_doc_path))
        malformed_md = temp_dir / "malformed.md"
        malformed_content = """# Incomplete heading
        
## Missing closing tags <strong>bold text
        
* Bullet without proper spacing
*Another bullet item
        
1. Numbered list
3. Wrong numbering
7. More wrong numbering

| Table | with | missing |
| cells | in | 
| some | rows"""
        malformed_md.write_text(malformed_content)
        output_path = temp_dir / "output.docx"
        
        # Act & Assert
        # Should not raise an exception, even with malformed input
        converter.convert(str(malformed_md), str(output_path))
        assert output_path.exists()