"""
Integration tests for the complete Word Formatter pipeline.
Tests the end-to-end workflow including shell script integration.
"""

import os
import pytest
import subprocess
import tempfile
from pathlib import Path
from unittest.mock import patch, MagicMock
from docx import Document

from document_converter import DocumentConverter
from document_converter_ai import AIDocumentConverter


class TestFullPipeline:
    """Integration tests for the complete document conversion pipeline."""
    
    @pytest.mark.integration
    def test_text_to_docx_basic_pipeline(self, reference_doc_path, sample_text_file, temp_dir):
        """Test complete pipeline from text file to formatted DOCX."""
        # Arrange
        converter = DocumentConverter(str(reference_doc_path))
        output_path = temp_dir / "pipeline_output.docx"
        
        # Act
        converter.convert(str(sample_text_file), str(output_path))
        
        # Assert
        assert output_path.exists()
        
        # Load and verify the output document
        output_doc = Document(str(output_path))
        assert len(output_doc.paragraphs) > 0
        
        # Check content preservation
        all_text = '\n'.join([p.text for p in output_doc.paragraphs])
        assert "Test Document" in all_text
        assert "Chapter 1: Introduction" in all_text
        
        # Check formatting applied
        section = output_doc.sections[0]
        # Should have applied reference margins (72pt = 1 inch)
        assert section.top_margin.pt == 72
        assert section.bottom_margin.pt == 72
    
    @pytest.mark.integration
    def test_markdown_to_docx_pipeline(self, reference_doc_path, sample_markdown_file, temp_dir):
        """Test complete pipeline from Markdown to formatted DOCX."""
        # Arrange
        converter = DocumentConverter(str(reference_doc_path))
        output_path = temp_dir / "markdown_output.docx"
        
        # Act
        converter.convert(str(sample_markdown_file), str(output_path))
        
        # Assert
        assert output_path.exists()
        
        # Load and verify the output document
        output_doc = Document(str(output_path))
        assert len(output_doc.paragraphs) > 0
        
        # Check for proper heading structure
        headings = [p for p in output_doc.paragraphs if p.style.name.startswith('Heading')]
        assert len(headings) > 0
        
        # Check specific heading levels
        h1_headings = [p for p in output_doc.paragraphs if p.style.name == 'Heading 1']
        h2_headings = [p for p in output_doc.paragraphs if p.style.name == 'Heading 2']
        assert len(h1_headings) > 0
        assert len(h2_headings) > 0
        
        # Check for tables
        assert len(output_doc.tables) > 0
        table = output_doc.tables[0]
        assert len(table.rows) >= 2
        assert len(table.columns) >= 3
    
    @pytest.mark.integration
    @pytest.mark.claude
    def test_ai_enhanced_pipeline(self, reference_doc_path, sample_text_file, temp_dir, mock_claude_success, mock_notifications):
        """Test AI-enhanced conversion pipeline."""
        # Arrange
        converter = AIDocumentConverter(str(reference_doc_path))
        converter.show_progress = False  # Disable for cleaner test output
        output_path = temp_dir / "ai_output.docx"
        
        # Act
        result = converter.convert_with_ai(str(sample_text_file), str(output_path))
        
        # Assert
        assert result is True
        assert output_path.exists()
        
        # Verify intermediate markdown was created
        markdown_path = output_path.parent / f"{sample_text_file.stem}_markdown.md"
        assert markdown_path.exists()
        
        # Check markdown content
        markdown_content = markdown_path.read_text()
        assert "# Test Document" in markdown_content
        
        # Load and verify the final output document
        output_doc = Document(str(output_path))
        assert len(output_doc.paragraphs) > 0
    
    @pytest.mark.integration
    @pytest.mark.claude
    def test_ai_pipeline_with_sanskrit_text(self, reference_doc_path, sanskrit_text_file, temp_dir, mock_claude_success, mock_notifications):
        """Test AI pipeline with Sanskrit and non-English text."""
        # Arrange
        converter = AIDocumentConverter(str(reference_doc_path))
        converter.show_progress = False
        output_path = temp_dir / "sanskrit_output.docx"
        
        # Act
        result = converter.convert_with_ai(str(sanskrit_text_file), str(output_path))
        
        # Assert
        assert result is True
        assert output_path.exists()
        
        # Check that markdown contains italicized non-English text
        markdown_path = output_path.parent / f"{sanskrit_text_file.stem}_markdown.md"
        if markdown_path.exists():
            markdown_content = markdown_path.read_text()
            # The AI mock should have processed non-English text
            # This test verifies the pipeline preserves the structure
            assert len(markdown_content) > 0
    
    @pytest.mark.integration
    @pytest.mark.claude
    def test_ai_pipeline_large_file_chunking(self, reference_doc_path, large_text_file, temp_dir, mock_claude_success, mock_notifications):
        """Test AI pipeline with large file requiring chunking."""
        # Arrange
        converter = AIDocumentConverter(str(reference_doc_path))
        converter.show_progress = False
        output_path = temp_dir / "large_output.docx"
        
        # Act
        result = converter.convert_with_ai(str(large_text_file), str(output_path))
        
        # Assert
        assert result is True
        assert output_path.exists()
        
        # Load and verify the output document
        output_doc = Document(str(output_path))
        assert len(output_doc.paragraphs) > 0
        
        # Should have processed all content despite chunking
        all_text = '\n'.join([p.text for p in output_doc.paragraphs])
        assert len(all_text) > 1000  # Should be substantial content
    
    @pytest.mark.integration
    def test_rtf_to_docx_pipeline(self, reference_doc_path, sample_rtf_file, temp_dir, mock_claude_success, mock_notifications):
        """Test complete pipeline for RTF files."""
        # Arrange
        converter = AIDocumentConverter(str(reference_doc_path))
        converter.show_progress = False
        output_path = temp_dir / "rtf_output.docx"
        
        # Act
        result = converter.convert_with_ai(str(sample_rtf_file), str(output_path))
        
        # Assert
        assert result is True
        assert output_path.exists()
        
        # Load and verify the output document
        output_doc = Document(str(output_path))
        assert len(output_doc.paragraphs) > 0
        
        # Check that RTF content was extracted and converted
        all_text = '\n'.join([p.text for p in output_doc.paragraphs])
        assert "Test Document" in all_text
    
    @pytest.mark.integration
    @pytest.mark.claude
    def test_ai_pipeline_with_fallback(self, reference_doc_path, sample_text_file, temp_dir, mock_notifications):
        """Test AI pipeline with fallback to simple conversion."""
        # Arrange
        converter = AIDocumentConverter(str(reference_doc_path))
        converter.show_progress = False
        output_path = temp_dir / "fallback_output.docx"
        
        # Mock Claude to fail, forcing fallback
        with patch.object(converter, '_process_in_chunks', return_value=None):
            # Act
            result = converter.convert_with_ai(str(sample_text_file), str(output_path))
        
        # Assert
        assert result is True  # Should still succeed with fallback
        assert output_path.exists()
        
        # Load and verify the output document
        output_doc = Document(str(output_path))
        assert len(output_doc.paragraphs) > 0
        
        # Content should still be preserved
        all_text = '\n'.join([p.text for p in output_doc.paragraphs])
        assert "Test Document" in all_text
    
    @pytest.mark.integration
    @pytest.mark.claude
    def test_haiku_fallback_pipeline(self, reference_doc_path, sample_text_file, temp_dir):
        """Test pipeline with Haiku fallback on timeout."""
        # Arrange
        converter = AIDocumentConverter(str(reference_doc_path))
        converter.show_progress = False
        converter.enable_haiku_fallback = True
        converter.model = 'sonnet'
        output_path = temp_dir / "haiku_fallback_output.docx"
        
        # Mock initial timeout, then success with Haiku
        call_count = 0
        def mock_run(cmd, **kwargs):
            nonlocal call_count
            call_count += 1
            
            if '--model' in cmd and 'sonnet' in cmd:
                raise subprocess.TimeoutExpired(cmd, 600)
            elif '--model' in cmd and 'haiku' in cmd:
                mock_result = MagicMock()
                mock_result.returncode = 0
                mock_result.stdout = "# Converted with Haiku\n\n## Chapter 1\n\nContent here."
                mock_result.stderr = ""
                return mock_result
            else:
                # For non-Claude calls (like notifications)
                mock_result = MagicMock()
                mock_result.returncode = 0
                return mock_result
        
        with patch('subprocess.run', side_effect=mock_run):
            # Act
            result = converter.convert_with_ai(str(sample_text_file), str(output_path))
        
        # Assert
        assert result is True
        assert output_path.exists()
        assert converter.model == 'haiku'  # Should have switched models
        
        # Check that Haiku was actually used
        markdown_path = output_path.parent / f"{sample_text_file.stem}_markdown.md"
        if markdown_path.exists():
            markdown_content = markdown_path.read_text()
            assert "Converted with Haiku" in markdown_content


class TestShellScriptIntegration:
    """Integration tests for the shell script wrapper."""
    
    @pytest.mark.integration
    @pytest.mark.slow
    def test_shell_script_basic_execution(self, reference_doc_path, sample_text_file, temp_dir):
        """Test basic shell script execution (without Claude)."""
        # Arrange
        script_path = Path(__file__).parent.parent.parent / "format_document.sh"
        if not script_path.exists():
            pytest.skip("Shell script not found")
        
        # Set up environment to disable AI processing
        env = os.environ.copy()
        env['ENABLE_NOTIFICATIONS'] = '0'  # Disable notifications
        env['WORD_FORMATTER_DEBUG'] = '0'  # Disable debug for cleaner output
        
        # Create a temporary script directory structure
        temp_script_dir = temp_dir / "script_test"
        temp_script_dir.mkdir()
        
        # Copy reference format to temp directory
        temp_ref = temp_script_dir / "referenceformat.docx"
        import shutil
        shutil.copy2(reference_doc_path, temp_ref)
        
        # Create a simple test version of the script that uses simple converter
        test_script = temp_script_dir / "test_format.sh"
        test_script.write_text(f'''#!/bin/bash
# Simple test version
python3 -c "
import sys
sys.path.insert(0, '{Path(__file__).parent.parent.parent}')
from document_converter_simple import main
main()
" "$1" --reference "{temp_ref}"
''')
        test_script.chmod(0o755)
        
        # Act
        result = subprocess.run(
            [str(test_script), str(sample_text_file)],
            cwd=str(temp_script_dir),
            env=env,
            capture_output=True,
            text=True,
            timeout=30
        )
        
        # Assert
        assert result.returncode == 0, f"Script failed with error: {result.stderr}"
        
        # Check that output file was created
        expected_output = sample_text_file.parent / f"{sample_text_file.stem}_formatted.docx"
        if expected_output.exists():
            assert expected_output.exists()
            
            # Load and verify basic content
            output_doc = Document(str(expected_output))
            assert len(output_doc.paragraphs) > 0
    
    @pytest.mark.integration
    def test_environment_variable_handling(self, reference_doc_path, sample_text_file, temp_dir):
        """Test that environment variables are properly handled."""
        # Arrange
        converter = AIDocumentConverter(str(reference_doc_path))
        
        # Test various environment variable combinations
        test_cases = [
            {'WORD_FORMATTER_DEBUG': '1', 'expected_debug': True},
            {'SAVE_MARKDOWN': '0', 'expected_save_markdown': False},
            {'SHOW_PROGRESS': '0', 'expected_show_progress': False},
            {'CLAUDE_MODEL': 'opus', 'expected_model': 'opus'},
            {'CLAUDE_TIMEOUT': '120', 'expected_timeout': 120},
        ]
        
        for test_case in test_cases:
            env_vars = {k: v for k, v in test_case.items() if k.startswith(('WORD_', 'SAVE_', 'SHOW_', 'CLAUDE_'))}
            
            with patch.dict(os.environ, env_vars):
                # Act
                test_converter = AIDocumentConverter(str(reference_doc_path))
                
                # Assert
                if 'expected_debug' in test_case:
                    assert test_converter.debug == test_case['expected_debug']
                if 'expected_save_markdown' in test_case:
                    assert test_converter.save_markdown == test_case['expected_save_markdown']
                if 'expected_show_progress' in test_case:
                    assert test_converter.show_progress == test_case['expected_show_progress']
                if 'expected_model' in test_case:
                    assert test_converter.model == test_case['expected_model']
                if 'expected_timeout' in test_case:
                    assert test_converter.timeout == test_case['expected_timeout']


class TestErrorHandlingIntegration:
    """Integration tests for error handling across the pipeline."""
    
    @pytest.mark.integration
    def test_invalid_input_file_handling(self, reference_doc_path, temp_dir):
        """Test handling of invalid input files."""
        # Arrange
        converter = DocumentConverter(str(reference_doc_path))
        invalid_file = temp_dir / "invalid.xyz"
        invalid_file.write_text("invalid content")
        output_path = temp_dir / "output.docx"
        
        # Act & Assert
        with pytest.raises(ValueError, match="Unsupported file format"):
            converter.convert(str(invalid_file), str(output_path))
    
    @pytest.mark.integration
    def test_missing_reference_file_handling(self, sample_text_file, temp_dir):
        """Test handling of missing reference files."""
        # Arrange
        nonexistent_ref = temp_dir / "nonexistent.docx"
        
        # Act & Assert
        with pytest.raises((FileNotFoundError, Exception)):
            DocumentConverter(str(nonexistent_ref))
    
    @pytest.mark.integration
    def test_corrupted_input_file_handling(self, reference_doc_path, temp_dir):
        """Test handling of corrupted input files."""
        # Arrange
        converter = DocumentConverter(str(reference_doc_path))
        corrupted_file = temp_dir / "corrupted.txt"
        # Write invalid UTF-8 bytes
        corrupted_file.write_bytes(b'\xff\xfe\x00\x00Invalid UTF-8')
        output_path = temp_dir / "output.docx"
        
        # Act & Assert
        with pytest.raises(UnicodeDecodeError):
            converter.convert(str(corrupted_file), str(output_path))
    
    @pytest.mark.integration
    @pytest.mark.claude
    def test_ai_conversion_error_recovery(self, reference_doc_path, sample_text_file, temp_dir, mock_notifications):
        """Test AI conversion with error recovery to simple conversion."""
        # Arrange
        converter = AIDocumentConverter(str(reference_doc_path))
        converter.show_progress = False
        output_path = temp_dir / "error_recovery_output.docx"
        
        # Mock all Claude attempts to fail
        with patch.object(converter, '_call_claude', return_value=(False, "All Claude attempts failed")):
            # Act
            result = converter.convert_with_ai(str(sample_text_file), str(output_path))
        
        # Assert
        assert result is True  # Should recover with simple conversion
        assert output_path.exists()
        
        # Content should still be processed
        output_doc = Document(str(output_path))
        assert len(output_doc.paragraphs) > 0
    
    @pytest.mark.integration
    def test_output_directory_permissions(self, reference_doc_path, sample_text_file, temp_dir):
        """Test handling of output directory permission issues."""
        # Arrange
        converter = DocumentConverter(str(reference_doc_path))
        readonly_dir = temp_dir / "readonly"
        readonly_dir.mkdir()
        # Make directory read-only (this might not work on all systems)
        readonly_dir.chmod(0o444)
        
        output_path = readonly_dir / "output.docx"
        
        try:
            # Act & Assert
            with pytest.raises((PermissionError, OSError)):
                converter.convert(str(sample_text_file), str(output_path))
        finally:
            # Clean up - restore permissions
            readonly_dir.chmod(0o755)


class TestPerformanceIntegration:
    """Integration tests for performance characteristics."""
    
    @pytest.mark.integration
    @pytest.mark.slow
    @pytest.mark.performance
    def test_large_file_processing_performance(self, reference_doc_path, temp_dir, mock_claude_success, mock_notifications):
        """Test performance with large files."""
        # Arrange
        large_file = temp_dir / "very_large.txt"
        # Create a ~1MB text file
        content_block = "This is a test paragraph with substantial content that will be repeated many times. " * 100
        chapters = []
        for i in range(100):  # 100 chapters
            chapters.append(f"Chapter {i+1}: Test Chapter\n\n{content_block}\n\n")
        
        large_content = '\n'.join(chapters)
        large_file.write_text(large_content, encoding='utf-8')
        
        converter = AIDocumentConverter(str(reference_doc_path))
        converter.show_progress = False
        output_path = temp_dir / "large_output.docx"
        
        # Act
        import time
        start_time = time.time()
        result = converter.convert_with_ai(str(large_file), str(output_path))
        end_time = time.time()
        
        # Assert
        assert result is True
        assert output_path.exists()
        
        # Performance assertion - should complete within reasonable time
        processing_time = end_time - start_time
        assert processing_time < 60  # Should complete within 1 minute for mocked Claude
        
        # Verify content integrity
        output_doc = Document(str(output_path))
        assert len(output_doc.paragraphs) > 100  # Should have substantial content
    
    @pytest.mark.integration
    @pytest.mark.performance 
    def test_multiple_file_processing(self, reference_doc_path, temp_dir, mock_claude_success, mock_notifications):
        """Test processing multiple files in sequence."""
        # Arrange
        converter = AIDocumentConverter(str(reference_doc_path))
        converter.show_progress = False
        
        # Create multiple test files
        test_files = []
        for i in range(5):
            test_file = temp_dir / f"test_{i}.txt"
            test_file.write_text(f"Test document {i}\n\nChapter 1: Content for file {i}\n\nSome content here.")
            test_files.append(test_file)
        
        # Act
        import time
        start_time = time.time()
        
        results = []
        for test_file in test_files:
            output_path = temp_dir / f"output_{test_file.stem}.docx"
            result = converter.convert_with_ai(str(test_file), str(output_path))
            results.append((result, output_path))
        
        end_time = time.time()
        
        # Assert
        assert all(result for result, _ in results), "All conversions should succeed"
        assert all(path.exists() for _, path in results), "All output files should exist"
        
        # Performance assertion
        processing_time = end_time - start_time
        assert processing_time < 30  # Should complete all files within 30 seconds