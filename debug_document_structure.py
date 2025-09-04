#!/usr/bin/env python3
"""Debug document structure to see what's actually in the Word file"""

from docx import Document
import sys

def debug_document(doc_path):
    doc = Document(doc_path)
    
    print(f"Analyzing document: {doc_path}")
    print("=" * 60)
    
    # Print first 20 paragraphs with their styles
    print("\nFirst 20 paragraphs:")
    print("-" * 40)
    
    for i, para in enumerate(doc.paragraphs[:20]):
        text = para.text.strip()
        style = para.style.name if para.style else "No style"
        
        # Check for page breaks
        page_break = ""
        if para.paragraph_format.page_break_before:
            page_break = " [PAGE BREAK BEFORE]"
        
        # Check for empty paragraphs that might contain page breaks
        for run in para.runs:
            if run._element.xml.find('<w:br w:type="page"/>') != -1:
                page_break += " [CONTAINS PAGE BREAK]"
        
        if text:
            print(f"{i:3}: [{style:20}]{page_break} {text[:60]}{'...' if len(text) > 60 else ''}")
        elif page_break:
            print(f"{i:3}: [Empty paragraph]{page_break}")
        else:
            print(f"{i:3}: [Empty paragraph]")
    
    print("\n" + "=" * 60)
    
    # Check for images
    print("\nChecking for images/shapes:")
    image_count = 0
    for para in doc.paragraphs:
        for run in para.runs:
            # Check for drawing elements using xml string search
            if '<w:drawing' in run._element.xml:
                image_count += 1
                print(f"  Found image in paragraph: {para.text[:30] if para.text else '[Empty]'}")
    print(f"Total images found: {image_count}")

if __name__ == "__main__":
    if len(sys.argv) > 1:
        debug_document(sys.argv[1])
    else:
        debug_document("test_samples/test_all_fixes_formatted.docx")