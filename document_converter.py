#!/usr/bin/env python3
"""
Smart Document Converter
Converts text, markdown, or Word documents to match a reference Word document's formatting.
"""

import argparse
import os
import sys
from pathlib import Path
from typing import Dict, List, Tuple, Optional
import warnings

# Import required libraries
try:
    from docx import Document
    from docx.shared import Pt, RGBColor, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
    from docx.enum.style import WD_STYLE_TYPE
    from docx.text.paragraph import Paragraph
    from docx.text.run import Run
    from docx.oxml import ns
except ImportError:
    print("Error: python-docx not installed. Run: pip install python-docx")
    sys.exit(1)

try:
    import markdown
    from markdown.extensions import tables, fenced_code, nl2br
except ImportError:
    print("Error: markdown not installed. Run: pip install markdown")
    sys.exit(1)

try:
    from bs4 import BeautifulSoup
except ImportError:
    print("Error: beautifulsoup4 not installed. Run: pip install beautifulsoup4")
    sys.exit(1)

try:
    from alias_resolver import resolve_path
except ImportError:
    # If alias_resolver is not available, create a dummy function
    def resolve_path(path):
        return str(path)

try:
    from config_loader import FormatterConfig
except ImportError:
    # If config_loader is not available, create a dummy
    FormatterConfig = None

try:
    from striprtf.striprtf import rtf_to_text
    RTF_SUPPORT = True
except ImportError:
    RTF_SUPPORT = False


class StyleExtractor:
    """Extracts formatting styles from a reference Word document."""
    
    def __init__(self, reference_doc_path: str):
        # Store the original path
        self.reference_doc_path = reference_doc_path
        # Resolve any aliases or symlinks
        resolved_path = resolve_path(reference_doc_path)
        self.reference_doc = Document(resolved_path)
        self.styles = self._extract_styles()
        self._log_extracted_styles()
    
    def _extract_styles(self) -> Dict:
        """Extract all relevant styles from the reference document."""
        styles = {
            'document': self._get_document_styles(),
            'paragraphs': self._get_paragraph_styles(),
            'headings': self._get_heading_styles(),
            'normal': self._get_normal_style(),
            'lists': self._get_list_styles()
        }
        return styles
    
    def _get_document_styles(self) -> Dict:
        """Extract document-level formatting."""
        doc_styles = {}
        
        # Extract page margins if available
        sections = self.reference_doc.sections
        if sections:
            section = sections[0]
            doc_styles['margins'] = {
                'top': section.top_margin,
                'bottom': section.bottom_margin,
                'left': section.left_margin,
                'right': section.right_margin
            }
            doc_styles['page_width'] = section.page_width
            doc_styles['page_height'] = section.page_height
        
        return doc_styles
    
    def _get_paragraph_styles(self) -> List[Dict]:
        """Extract paragraph styles from the document."""
        paragraph_styles = []
        seen_styles = set()
        
        for para in self.reference_doc.paragraphs:
            if para.text.strip() and para.style.name not in seen_styles:
                seen_styles.add(para.style.name)
                
                style = {
                    'alignment': para.alignment,
                    'left_indent': para.paragraph_format.left_indent,
                    'right_indent': para.paragraph_format.right_indent,
                    'first_line_indent': para.paragraph_format.first_line_indent,
                    'space_before': para.paragraph_format.space_before,
                    'space_after': para.paragraph_format.space_after,
                    'line_spacing': para.paragraph_format.line_spacing,
                    'style_name': para.style.name if para.style else None
                }
                
                # Extract run formatting from all runs to get most common style
                if para.runs:
                    # Get the most common font properties
                    fonts = []
                    for run in para.runs:
                        if run.text.strip():
                            fonts.append({
                                'name': run.font.name,
                                'size': run.font.size,
                                'bold': run.font.bold,
                                'italic': run.font.italic,
                                'underline': run.font.underline,
                                'color': run.font.color.rgb if run.font.color else None
                            })
                    
                    # Use the first non-empty run's style
                    if fonts:
                        style['font'] = fonts[0]
                
                paragraph_styles.append(style)
                
                # Stop after collecting enough unique styles
                if len(paragraph_styles) >= 20:
                    break
        
        return paragraph_styles
    
    def _get_heading_styles(self) -> Dict:
        """Extract heading styles from the document."""
        heading_styles = {}
        
        for style in self.reference_doc.styles:
            if style.type == WD_STYLE_TYPE.PARAGRAPH and 'Heading' in style.name:
                heading_styles[style.name] = {
                    'font': {
                        'name': style.font.name,
                        'size': style.font.size,
                        'bold': style.font.bold,
                        'italic': style.font.italic,
                        'color': style.font.color.rgb if style.font.color else None
                    },
                    'paragraph': {
                        'alignment': style.paragraph_format.alignment,
                        'space_before': style.paragraph_format.space_before,
                        'space_after': style.paragraph_format.space_after,
                        'keep_with_next': style.paragraph_format.keep_with_next,
                        'page_break_before': style.paragraph_format.page_break_before,
                        'left_indent': style.paragraph_format.left_indent,
                        'right_indent': style.paragraph_format.right_indent,
                        'first_line_indent': style.paragraph_format.first_line_indent
                    }
                }
        
        # Also extract Title style if present
        for style in self.reference_doc.styles:
            if style.type == WD_STYLE_TYPE.PARAGRAPH and style.name == 'Title':
                heading_styles['Title'] = {
                    'font': {
                        'name': style.font.name,
                        'size': style.font.size,
                        'bold': style.font.bold,
                        'italic': style.font.italic,
                        'color': style.font.color.rgb if style.font.color else None
                    },
                    'paragraph': {
                        'alignment': style.paragraph_format.alignment,
                        'space_before': style.paragraph_format.space_before,
                        'space_after': style.paragraph_format.space_after,
                        'page_break_before': style.paragraph_format.page_break_before
                    }
                }
                
        return heading_styles
    
    def _get_normal_style(self) -> Dict:
        """Extract the normal/body text style."""
        normal_style = {}
        
        # First try to get from style definition
        for style in self.reference_doc.styles:
            if style.name == 'Normal':
                normal_style = {
                    'font': {
                        'name': style.font.name,
                        'size': style.font.size,
                        'bold': style.font.bold,
                        'italic': style.font.italic
                    },
                    'paragraph': {
                        'alignment': style.paragraph_format.alignment,
                        'space_before': style.paragraph_format.space_before,
                        'space_after': style.paragraph_format.space_after,
                        'line_spacing': style.paragraph_format.line_spacing
                    }
                }
                break
        
        # Override with actual font used in Normal style paragraphs
        for para in self.reference_doc.paragraphs:
            if para.style.name == 'Normal' and para.runs and para.text.strip():
                for run in para.runs:
                    if run.text.strip():
                        # Use the actual font from the document
                        normal_style['font'] = {
                            'name': run.font.name or normal_style.get('font', {}).get('name'),
                            'size': run.font.size or normal_style.get('font', {}).get('size'),
                            'bold': run.font.bold if run.font.bold is not None else normal_style.get('font', {}).get('bold'),
                            'italic': run.font.italic if run.font.italic is not None else normal_style.get('font', {}).get('italic')
                        }
                        break
                break
        
        return normal_style
    
    def _get_list_styles(self) -> Dict:
        """Extract list formatting styles."""
        list_styles = {
            'bullet': {'indent': Pt(36), 'font': None},
            'numbered': {'indent': Pt(36), 'font': None}
        }
        
        # Try to extract actual list styles from reference document
        for style in self.reference_doc.styles:
            if style.name == 'List Bullet' and hasattr(style, 'font'):
                list_styles['bullet']['font'] = {
                    'name': style.font.name,
                    'size': style.font.size
                }
            elif style.name == 'List Number' and hasattr(style, 'font'):
                list_styles['numbered']['font'] = {
                    'name': style.font.name,
                    'size': style.font.size
                }
        
        return list_styles
    
    def export_styles_to_file(self, output_path: str):
        """Export all extracted styles to a comprehensive JSON file."""
        import json
        import datetime
        
        # Analyze all styles in the document
        all_styles = self._get_all_document_styles()
        
        export_data = {
            "metadata": {
                "reference_document": str(self.reference_doc_path),
                "extraction_date": datetime.datetime.now().isoformat(),
                "document_name": self.reference_doc.core_properties.title or "Untitled",
                "author": self.reference_doc.core_properties.author or "Unknown"
            },
            "extracted_styles": self.styles,  # Already extracted styles
            "all_styles": all_styles,  # Comprehensive style list
            "font_usage": self._analyze_font_usage(),
            "paragraph_samples": self._get_paragraph_samples(),
            "style_hierarchy": self._get_style_hierarchy()
        }
        
        # Convert any non-serializable objects
        export_data = self._make_json_serializable(export_data)
        
        with open(output_path, 'w', encoding='utf-8') as f:
            json.dump(export_data, f, indent=2, ensure_ascii=False)
        
        print(f"\n✅ Style details exported to: {output_path}")
        print(f"   Total styles found: {len(all_styles)}")
        print(f"   Headings extracted: {len(self.styles.get('headings', {}))}")
        
    def _get_all_document_styles(self):
        """Get ALL styles defined in the document."""
        all_styles = {}
        
        for style in self.reference_doc.styles:
            try:
                style_info = {
                    'name': style.name,
                    'type': str(style.type),
                    'base_style': style.base_style.name if hasattr(style, 'base_style') and style.base_style else None,
                    'is_builtin': style.builtin if hasattr(style, 'builtin') else None,
                    'priority': style.priority if hasattr(style, 'priority') else None
                }
                
                # For paragraph styles, extract detailed formatting
                if style.type == WD_STYLE_TYPE.PARAGRAPH:
                    style_info.update({
                        'font': {
                            'name': style.font.name,
                            'size': style.font.size,
                            'bold': style.font.bold,
                            'italic': style.font.italic,
                            'underline': style.font.underline,
                            'color': style.font.color.rgb if style.font.color else None,
                            'highlight_color': style.font.highlight_color if hasattr(style.font, 'highlight_color') else None
                        },
                        'paragraph': {
                            'alignment': style.paragraph_format.alignment,
                            'left_indent': style.paragraph_format.left_indent,
                            'right_indent': style.paragraph_format.right_indent,
                            'first_line_indent': style.paragraph_format.first_line_indent,
                            'space_before': style.paragraph_format.space_before,
                            'space_after': style.paragraph_format.space_after,
                            'line_spacing': style.paragraph_format.line_spacing,
                            'line_spacing_rule': style.paragraph_format.line_spacing_rule if hasattr(style.paragraph_format, 'line_spacing_rule') else None,
                            'keep_with_next': style.paragraph_format.keep_with_next,
                            'keep_together': style.paragraph_format.keep_together,
                            'page_break_before': style.paragraph_format.page_break_before,
                            'widow_control': style.paragraph_format.widow_control
                        }
                    })
                    
                all_styles[style.name] = style_info
            except Exception as e:
                all_styles[style.name] = {'error': str(e)}
                
        return all_styles
    
    def _analyze_font_usage(self):
        """Analyze font usage throughout the document."""
        font_usage = {}
        
        for para in self.reference_doc.paragraphs:
            if para.runs and para.text.strip():
                for run in para.runs:
                    if run.text.strip() and run.font.name:
                        font_name = run.font.name
                        if font_name not in font_usage:
                            font_usage[font_name] = {
                                'count': 0,
                                'sizes': [],
                                'colors': [],
                                'styles': {'bold': False, 'italic': False, 'underline': False}
                            }
                        font_usage[font_name]['count'] += 1
                        
                        if run.font.size:
                            size_pt = run.font.size.pt if hasattr(run.font.size, 'pt') else run.font.size / 12700
                            if size_pt not in font_usage[font_name]['sizes']:
                                font_usage[font_name]['sizes'].append(size_pt)
                                
                        if run.font.color and run.font.color.rgb:
                            color = str(run.font.color.rgb)
                            if color not in font_usage[font_name]['colors']:
                                font_usage[font_name]['colors'].append(color)
                                
                        if run.font.bold:
                            font_usage[font_name]['styles']['bold'] = True
                        if run.font.italic:
                            font_usage[font_name]['styles']['italic'] = True
                        if run.font.underline:
                            font_usage[font_name]['styles']['underline'] = True
        
        # Sort sizes for each font
        for font_name in font_usage:
            font_usage[font_name]['sizes'].sort()
            
        return font_usage
    
    def _get_paragraph_samples(self):
        """Get sample paragraphs with their complete style information."""
        samples = []
        style_count = {}
        
        for i, para in enumerate(self.reference_doc.paragraphs[:100]):  # Sample first 100 paragraphs
            if para.text.strip():
                style_name = para.style.name if para.style else 'None'
                
                # Track style usage count
                style_count[style_name] = style_count.get(style_name, 0) + 1
                
                # Only include first few examples of each style
                if style_count[style_name] <= 3:
                    sample = {
                        'index': i,
                        'text': para.text[:150] + '...' if len(para.text) > 150 else para.text,
                        'style': style_name,
                        'alignment': str(para.alignment) if para.alignment else 'Default',
                        'indentation': {
                            'left': para.paragraph_format.left_indent.pt if para.paragraph_format.left_indent else 0,
                            'right': para.paragraph_format.right_indent.pt if para.paragraph_format.right_indent else 0,
                            'first_line': para.paragraph_format.first_line_indent.pt if para.paragraph_format.first_line_indent else 0
                        },
                        'spacing': {
                            'before': para.paragraph_format.space_before.pt if para.paragraph_format.space_before else 0,
                            'after': para.paragraph_format.space_after.pt if para.paragraph_format.space_after else 0,
                            'line_spacing': str(para.paragraph_format.line_spacing) if para.paragraph_format.line_spacing else 'Default'
                        },
                        'runs': []
                    }
                    
                    # Add run details
                    for run in para.runs[:5]:  # First 5 runs
                        if run.text.strip():
                            sample['runs'].append({
                                'text': run.text[:50] + '...' if len(run.text) > 50 else run.text,
                                'font': run.font.name,
                                'size': run.font.size.pt if run.font.size and hasattr(run.font.size, 'pt') else None,
                                'bold': run.font.bold,
                                'italic': run.font.italic,
                                'color': str(run.font.color.rgb) if run.font.color and run.font.color.rgb else None
                            })
                    
                    samples.append(sample)
        
        return samples
    
    def _get_style_hierarchy(self):
        """Get the style hierarchy showing which styles are based on others."""
        hierarchy = {}
        
        for style in self.reference_doc.styles:
            if hasattr(style, 'base_style') and style.base_style:
                base_name = style.base_style.name
                if base_name not in hierarchy:
                    hierarchy[base_name] = []
                hierarchy[base_name].append(style.name)
                
        return hierarchy
    
    def _make_json_serializable(self, obj):
        """Convert non-serializable objects to serializable format."""
        if isinstance(obj, dict):
            return {k: self._make_json_serializable(v) for k, v in obj.items()}
        elif isinstance(obj, list):
            return [self._make_json_serializable(v) for v in obj]
        elif isinstance(obj, set):
            return list(obj)
        elif hasattr(obj, 'pt'):
            return f"{obj.pt}pt"
        elif isinstance(obj, (int, float)) and obj > 10000:
            # Likely a Word internal unit, convert to points
            return f"{obj/12700:.1f}pt"
        elif obj is None:
            return None
        else:
            return str(obj)
            
    def _log_extracted_styles(self):
        """Log all extracted styles for debugging."""
        import json
        import os
        
        # Only log if debug mode is enabled
        if os.environ.get('WORD_FORMATTER_DEBUG', '0') != '1':
            return
            
        print("\n" + "="*60)
        print("REFERENCE DOCUMENT STYLES EXTRACTED")
        print("="*60)
        
        # Show actual fonts used in the document
        print("\n🔍 ACTUAL FONTS DETECTED IN DOCUMENT:")
        font_usage = {}
        for para in self.reference_doc.paragraphs[:20]:  # Sample first 20 paragraphs
            if para.runs and para.text.strip():
                for run in para.runs:
                    if run.text.strip() and run.font.name:
                        font_name = run.font.name
                        if font_name not in font_usage:
                            font_usage[font_name] = {
                                'count': 0,
                                'sizes': set(),
                                'styles': {'bold': False, 'italic': False}
                            }
                        font_usage[font_name]['count'] += 1
                        if run.font.size:
                            font_usage[font_name]['sizes'].add(run.font.size)
                        if run.font.bold:
                            font_usage[font_name]['styles']['bold'] = True
                        if run.font.italic:
                            font_usage[font_name]['styles']['italic'] = True
        
        for font_name, usage in font_usage.items():
            print(f"  {font_name}:")
            print(f"    Used {usage['count']} times")
            if usage['sizes']:
                sizes_pt = [f"{s/12700:.1f}pt" for s in sorted(usage['sizes'])]
                print(f"    Sizes: {', '.join(sizes_pt)}")
            if usage['styles']['bold'] or usage['styles']['italic']:
                styles = []
                if usage['styles']['bold']:
                    styles.append('Bold')
                if usage['styles']['italic']:
                    styles.append('Italic')
                print(f"    Styles: {', '.join(styles)}")
        
        print("\n" + "="*60)
        
        # Document-level styles
        if 'document' in self.styles:
            print("\n📄 DOCUMENT SETTINGS:")
            doc = self.styles['document']
            if 'margins' in doc:
                margins = doc['margins']
                print(f"  Margins:")
                print(f"    Top: {margins.get('top', 'Not set')}")
                print(f"    Bottom: {margins.get('bottom', 'Not set')}")
                print(f"    Left: {margins.get('left', 'Not set')}")
                print(f"    Right: {margins.get('right', 'Not set')}")
            if 'page_width' in doc:
                print(f"  Page Width: {doc['page_width']}")
            if 'page_height' in doc:
                print(f"  Page Height: {doc['page_height']}")
        
        # Normal style
        if 'normal' in self.styles and self.styles['normal']:
            print("\n📝 NORMAL (BODY) STYLE:")
            normal = self.styles['normal']
            if 'font' in normal:
                font = normal['font']
                print(f"  Font:")
                print(f"    Name: {font.get('name', 'Default')}")
                print(f"    Size: {font.get('size', 'Default')}")
                print(f"    Bold: {font.get('bold', False)}")
                print(f"    Italic: {font.get('italic', False)}")
            if 'paragraph' in normal:
                para = normal['paragraph']
                print(f"  Paragraph:")
                print(f"    Alignment: {para.get('alignment', 'Default')}")
                print(f"    Space Before: {para.get('space_before', 0)}")
                print(f"    Space After: {para.get('space_after', 0)}")
                print(f"    Line Spacing: {para.get('line_spacing', 'Default')}")
        
        # Heading styles
        if 'headings' in self.styles:
            print("\n📌 HEADING STYLES:")
            for heading_name, heading_style in self.styles['headings'].items():
                print(f"\n  {heading_name}:")
                if 'font' in heading_style:
                    font = heading_style['font']
                    print(f"    Font: {font.get('name', 'Default')}")
                    print(f"    Size: {font.get('size', 'Default')}")
                    print(f"    Bold: {font.get('bold', False)}")
                    print(f"    Color: {font.get('color', 'Default')}")
                if 'paragraph' in heading_style:
                    para = heading_style['paragraph']
                    print(f"    Space Before: {para.get('space_before', 0)}")
                    print(f"    Space After: {para.get('space_after', 0)}")
                    print(f"    Keep with Next: {para.get('keep_with_next', False)}")
                    print(f"    Page Break Before: {para.get('page_break_before', False)}")
        
        # List styles
        if 'lists' in self.styles:
            print("\n📋 LIST STYLES:")
            lists = self.styles['lists']
            if 'bullet' in lists:
                print(f"  Bullet List:")
                print(f"    Indent: {lists['bullet'].get('indent', 'Default')}")
                if lists['bullet'].get('font'):
                    font = lists['bullet']['font']
                    print(f"    Font: {font.get('name', 'Default')}")
                    print(f"    Size: {font.get('size', 'Default')}")
            if 'numbered' in lists:
                print(f"  Numbered List:")
                print(f"    Indent: {lists['numbered'].get('indent', 'Default')}")
                if lists['numbered'].get('font'):
                    font = lists['numbered']['font']
                    print(f"    Font: {font.get('name', 'Default')}")
                    print(f"    Size: {font.get('size', 'Default')}")
        
        # Paragraph styles summary
        if 'paragraphs' in self.styles:
            print(f"\n📊 PARAGRAPH STYLES COLLECTED: {len(self.styles['paragraphs'])} unique styles")
            style_names = set()
            for para_style in self.styles['paragraphs']:
                if para_style.get('style_name'):
                    style_names.add(para_style['style_name'])
            if style_names:
                print(f"  Style names found: {', '.join(sorted(style_names))}")
        
        print("\n" + "="*60)
        print()


class DocumentConverter:
    """Converts various document formats to Word with reference styling."""
    
    def __init__(self, reference_doc_path: str, config_path: Optional[str] = None):
        # Store reference path for template use
        self.reference_doc_path = reference_doc_path
        resolved_path = resolve_path(reference_doc_path)
        
        # Check if we should use formatter config
        # For now, always load config for behavioral rules (page breaks, chapter detection)
        # but only apply style overrides if USE_FORMATTER_CONFIG=1
        self.use_config = os.environ.get('USE_FORMATTER_CONFIG', '0') == '1'
        self.always_use_behavioral_config = True  # Always use config for page breaks, chapter detection
        
        # Use reference document as template to preserve all styles
        print(f"Using template approach with reference: {resolved_path}")
        self.output_doc = Document(resolved_path)
        
        # Clear all content from template while preserving styles
        self._clear_template_content()
        
        # Always load config for behavioral rules
        if FormatterConfig:
            self.config = FormatterConfig(config_path)
        else:
            self.config = None
        
        # Only extract styles if config is enabled (for style overrides)
        if self.use_config and self.config:
            self.style_extractor = StyleExtractor(reference_doc_path)
        else:
            self.style_extractor = None
        
        # Track chapter state for separator placement
        self.current_chapter_started = False
        self.current_chapter_name = ""
        self.paragraphs_since_chapter = 0
        self.current_chapter_elements = []
        
        # Track whether we've processed the first H1 as title
        self.first_h1_as_title = False
        
        # Track hierarchical list context
        self.in_hierarchical_list = False
        self.current_list_heading_level = None
        
        # Track special sections for page breaks
        self.special_sections = {
            'title': False,
            'dedication': False,
            'contents': False,
            'preface': False,
            'foreword': False
        }
    
    def _clear_template_content(self):
        """Clear all content from the template document while preserving styles."""
        # Clear all paragraphs from the document
        for paragraph in list(self.output_doc.paragraphs):
            p_element = paragraph._element
            p_element.getparent().remove(p_element)
        
        # Clear all tables if any
        for table in list(self.output_doc.tables):
            t_element = table._element
            t_element.getparent().remove(t_element)
        
        print("Template content cleared, styles preserved")
    
    def _is_special_section(self, text: str) -> Optional[str]:
        """Check if text represents a special section that needs its own page."""
        text_lower = text.lower().strip()
        
        # No automatic title detection - first H1 is handled separately
        
        if 'dedication' in text_lower or 'dedicated to' in text_lower or 'in memory of' in text_lower:
            return 'dedication'
        
        if text_lower in ['contents', 'table of contents', 'toc'] or 'table of contents' in text_lower:
            return 'contents'
        
        if text_lower == 'preface':
            return 'preface'
        
        if text_lower == 'foreword':
            return 'foreword'
        
        return None
    
    def _handle_special_section_page_break(self, section_type: str):
        """Mark special section as processed. Page breaks will be added before next major section."""
        if section_type and not self.special_sections.get(section_type, False):
            self.special_sections[section_type] = True
            if os.environ.get('WORD_FORMATTER_DEBUG', '0') == '1':
                print(f"DEBUG: Processing special section: {section_type} (page break deferred)")
            else:
                print(f"Processing special section: {section_type}")
            
            # Don't add immediate page breaks after dedication or contents
            # They should flow together until we hit the next major section
            # Only add immediate page break after title
            if self.config and len(self.output_doc.paragraphs) > 0:
                should_add_pagebreak = False
                if section_type == 'title' and self.config.should_add_page_break_after_title():
                    should_add_pagebreak = True
                
                if should_add_pagebreak:
                    self.output_doc.add_page_break()
                    if os.environ.get('WORD_FORMATTER_DEBUG', '0') == '1':
                        print(f"DEBUG: Added immediate page break AFTER {section_type} section")
                    else:
                        print(f"Added page break after {section_type}")
    
    def _handle_chapter_end(self):
        """Handle formatting at the end of a chapter."""
        # Check if we need to apply closing content formatting
        if self.current_chapter_elements and self.config:
            closing_settings = self.config.get_chapter_closing_settings()
            if closing_settings.get('enabled'):
                # Look for closing content patterns in the last few elements
                for i in range(max(0, len(self.current_chapter_elements) - 5), len(self.current_chapter_elements)):
                    elem_info = self.current_chapter_elements[i]
                    if self._is_closing_content(elem_info['text'], closing_settings):
                        # Apply subtle emphasis to closing content
                        para = elem_info['paragraph']
                        if para and para.runs:
                            self._apply_subtle_emphasis(para, closing_settings)
        
        # Chapter separator will be added when next chapter starts (not here)
        if os.environ.get('WORD_FORMATTER_DEBUG', '0') == '1':
            print("DEBUG: Chapter end processing completed (separator will be added when next chapter starts)")
    
    def _is_chapter_opening_content(self, text: str, position: int) -> bool:
        """Check if this paragraph is chapter opening content (verse/quote)."""
        if position > 3:  # Only check first few paragraphs after chapter heading
            return False
            
        opening_settings = self.config.get_chapter_opening_settings()
        if not opening_settings.get('enabled'):
            return False
            
        patterns = opening_settings.get('detect_patterns', [])
        text_lower = text.lower()
        
        # Check for verse/quote patterns
        for pattern in patterns:
            if pattern in text_lower:
                return True
                
        # Also check for common quote indicators
        if text.strip().startswith(('"', '“', '‘', '—')) or text.strip().endswith(('"', '”', '’')):
            return True
            
        return False
    
    def _is_closing_content(self, text: str, settings: Dict) -> bool:
        """Check if this paragraph is closing content (story/poem)."""
        patterns = settings.get('detect_patterns', [])
        text_lower = text.lower()
        
        for pattern in patterns:
            if pattern in text_lower:
                return True
                
        return False
    
    def _get_reference_font_for_style(self, style_name: str) -> Optional[Dict]:
        """Extract font information for a given style from the reference document."""
        if not self.style_extractor or not hasattr(self.style_extractor, 'styles'):
            return None
        
        # Look for paragraphs using this style in the extracted styles
        paragraphs = self.style_extractor.styles.get('paragraphs', [])
        for para in paragraphs:
            if para.get('style_name') == style_name:
                # Look for the actual font used in runs
                if 'runs' in para:
                    for run in para['runs']:
                        font = run.get('font')
                        if font and font != 'null' and font != 'None':
                            return {
                                'name': font,
                                'size': run.get('size'),
                                'bold': run.get('bold'),
                                'italic': run.get('italic'),
                                'color': run.get('color')
                            }
                            
        # If no runs found, check if style has font definition
        headings = self.style_extractor.styles.get('headings', {})
        if style_name in headings:
            style_info = headings[style_name]
            font_info = style_info.get('font', {})
            if font_info and font_info.get('name'):
                return {
                    'name': font_info.get('name'),
                    'size': font_info.get('size'),
                    'bold': font_info.get('bold'),
                    'italic': font_info.get('italic'),
                    'color': font_info.get('color')
                }
        
        return None
    
    def _apply_subtle_emphasis(self, paragraph, settings: Dict):
        """Apply subtle emphasis formatting to a paragraph."""
        if os.environ.get('WORD_FORMATTER_DEBUG', '0') == '1':
            print(f"DEBUG: Applying subtle emphasis with settings: {settings}")
            print(f"DEBUG: Paragraph has {len(paragraph.runs)} runs")
        
        # Apply to all runs in the paragraph
        for i, run in enumerate(paragraph.runs):
            if os.environ.get('WORD_FORMATTER_DEBUG', '0') == '1':
                print(f"DEBUG: Run {i}: '{run.text[:30]}...', applying italic={settings.get('italic')}")
            
            if settings.get('italic'):
                run.italic = True
                
            if settings.get('color'):
                run.font.color.rgb = RGBColor.from_string(settings['color'])
                
            if settings.get('font_size_reduction'):
                current_size = run.font.size
                if current_size:
                    run.font.size = Pt(current_size.pt - settings['font_size_reduction'])
        
        # Apply paragraph spacing
        if settings.get('spacing_before'):
            paragraph.paragraph_format.space_before = Pt(settings['spacing_before'])
        if settings.get('spacing_after'):
            paragraph.paragraph_format.space_after = Pt(settings['spacing_after'])
    
    def _add_chapter_separator(self, separator_settings):
        """Add a chapter separator based on settings."""
        sep_para = self.output_doc.add_paragraph()
        sep_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        sep_para.paragraph_format.space_before = Pt(separator_settings.get('spacing_before', 12))
        sep_para.paragraph_format.space_after = Pt(separator_settings.get('spacing_after', 12))
        
        # Use image if available, otherwise use text symbol
        if separator_settings.get('use_image') and separator_settings.get('image_data'):
            import io
            image_data = separator_settings['image_data']
            width = separator_settings.get('image_width', 0.5)
            height = separator_settings.get('image_height', 0.5)
            sep_para.add_run().add_picture(io.BytesIO(image_data), 
                                         width=Inches(width), 
                                         height=Inches(height))
        else:
            # Use text symbol
            symbol = separator_settings.get('symbol', '❦')
            run = sep_para.add_run(symbol)
            run.font.size = Pt(separator_settings.get('font_size', 14))
        
    def convert(self, input_path: str, output_path: str):
        """Convert input document to styled Word document."""
        # Resolve any aliases or symlinks
        input_path = Path(resolve_path(input_path))

        # Determine file type and read content
        if input_path.suffix.lower() == '.txt':
            content = self._read_text_file(input_path)
            self._process_text_content(content)
        elif input_path.suffix.lower() == '.rtf':
            if not RTF_SUPPORT:
                raise ValueError("RTF support requires striprtf library. Install with: pip install striprtf")
            # Read RTF and convert to plain text
            print(f"Converting RTF to plain text...")
            with open(input_path, 'r', encoding='utf-8') as f:
                rtf_content = f.read()
            content = rtf_to_text(rtf_content)
            print(f"Extracted {len(content)} characters from RTF")
            # Process as text content (striprtf gives us plain text)
            self._process_text_content(content)
        elif input_path.suffix.lower() in ['.md', '.markdown']:
            content = self._read_markdown_file(input_path)
            self._process_markdown_content(content)
        elif input_path.suffix.lower() in ['.docx', '.doc']:
            self._process_word_document(input_path)
        else:
            raise ValueError(f"Unsupported file format: {input_path.suffix}")
        
        # Apply reference document styles
        self._apply_reference_styles()
        
        # Save the output document
        self.output_doc.save(output_path)
        print(f"Document converted successfully: {output_path}")
    
    def _read_text_file(self, file_path: Path) -> str:
        """Read plain text file."""
        with open(file_path, 'r', encoding='utf-8') as f:
            return f.read()
    
    def _read_markdown_file(self, file_path: Path) -> str:
        """Read markdown file."""
        with open(file_path, 'r', encoding='utf-8') as f:
            return f.read()
    
    def _process_text_content(self, content: str):
        """Process plain text content."""
        paragraphs = content.split('\n\n')
        
        for para_text in paragraphs:
            if para_text.strip():
                para = self.output_doc.add_paragraph(para_text.strip())
    
    def _process_markdown_content(self, content: str):
        """Process markdown content to Word format."""
        # First, handle any explicit page break markers
        # Common markdown page break representations
        page_break_markers = [
            '\\newpage',
            '\\pagebreak',
            '<div style="page-break-after: always;"></div>',
            '<div style="page-break-before: always;"></div>',
            '---pagebreak---',
            '<!-- pagebreak -->',
            '<pb>',
            '</pb>'
        ]
        
        # Replace page break markers with a unique placeholder
        for marker in page_break_markers:
            content = content.replace(marker, '<!--PAGEBREAK-->')
        
        # No special title processing needed - first H1 will be treated as title
        
        # Convert markdown to HTML
        extensions = ['tables', 'fenced_code', 'nl2br']
        html = markdown.markdown(content, extensions=extensions)
        if os.environ.get('WORD_FORMATTER_DEBUG', '0') == '1' and os.environ.get('SHOW_HTML', '0') == '1':
            print(f"DEBUG: Generated HTML from markdown:")
            print(html[:800] + "..." if len(html) > 800 else html)
        
        # Parse HTML
        soup = BeautifulSoup(html, 'html.parser')
        
        # No title placeholder processing needed
        
        # Track processed elements to avoid duplicates
        processed_elements = set()
        
        # Get all elements
        all_elements = soup.find_all(['h0', 'h1', 'h2', 'h3', 'h4', 'h5', 'h6', 'p', 'blockquote', 'ul', 'ol', 'table', 'pre'])
        if os.environ.get('WORD_FORMATTER_DEBUG', '0') == '1':
            print(f"DEBUG: Found {len(all_elements)} elements to process:")
            for i, elem in enumerate(all_elements):
                content = elem.get_text().strip()[:50]
                print(f"  {i}: <{elem.name}> '{content}...'")
        
        # Track when we're at the end of a chapter for separator placement
        self.current_chapter_started = False
        
        # Process each element
        for idx, element in enumerate(all_elements):
            # Check if the next element is a new chapter to add separator before it
            next_is_chapter = False
            if idx + 1 < len(all_elements):
                next_elem = all_elements[idx + 1]
                if next_elem.name in ['h0', 'h1', 'h2', 'h3', 'h4', 'h5', 'h6']:
                    # h0 after the first one is a chapter
                    if next_elem.name == 'h0' and self.special_sections['title']:
                        next_is_chapter = True
                    else:
                        level = 0 if next_elem.name == 'h0' else int(next_elem.name[1])
                        heading_text = next_elem.get_text().strip()
                        if self.config and (level == 1 or self.config.is_chapter_keyword(heading_text)):
                            next_is_chapter = True
            
            # Add separator at end of chapter if configured for "after" position
            if self.current_chapter_started and next_is_chapter and self.config:
                self._handle_chapter_end()
            
            self._process_html_element(element, processed_elements)
            
            # Track elements for chapter formatting
            if self.current_chapter_started:
                self.paragraphs_since_chapter += 1
            
        # Add separator at the very end only if this is the final chapter
        if self.current_chapter_started and self.config:
            # Add separator at end of final chapter
            separator_settings = self.config.get_chapter_separator()
            if separator_settings and separator_settings.get('enabled') and separator_settings.get('position') == 'after':
                self._add_chapter_separator(separator_settings)
                if os.environ.get('WORD_FORMATTER_DEBUG', '0') == '1':
                    print(f"DEBUG: Added final chapter separator at end of document for '{self.current_chapter_name}'")
                else:
                    print(f"Added separator at end of final chapter: {self.current_chapter_name}")
    
    def _process_html_element(self, element, processed_elements):
        """Process individual HTML elements."""
        # Skip if already processed (to avoid duplicates in blockquotes)
        if id(element) in processed_elements:
            return
        if element.name in ['h1', 'h2', 'h3', 'h4', 'h5', 'h6']:
            level = int(element.name[1])
            
            # Handle first H1 as Title
            if element.name == 'h1' and not self.first_h1_as_title:
                use_title_style = True
                level = 0  # Treat as title level
                self.first_h1_as_title = True
                if os.environ.get('WORD_FORMATTER_DEBUG', '0') == '1':
                    print(f"DEBUG: First H1 converted to document title: '{element.get_text().strip()}'")
            else:
                use_title_style = False
            
            heading_text = element.get_text().strip()
            
            # Determine the Word heading level
            word_heading_level = level
            
            # Use configuration if available for behavioral rules
            if self.config and self.always_use_behavioral_config:
                # Check if this is a section (use Heading 1)
                if self.config.is_section_keyword(heading_text):
                    word_heading_level = 1
                    # Check for page break setting
                    if self.config.should_apply_page_break("section") and len(self.output_doc.paragraphs) > 0:
                        self.output_doc.add_page_break()
                # Check if this is a chapter (always use Heading 2)
                elif level == 1 or self.config.is_chapter_keyword(heading_text):
                    word_heading_level = 2
                    
                    # Debug logging
                    if os.environ.get('WORD_FORMATTER_DEBUG', '0') == '1':
                        print(f"DEBUG: Chapter detected: '{heading_text}' (level={level}, keyword_match={self.config.is_chapter_keyword(heading_text)})")
                    
                    # Add separator at end of previous chapter if we were in one
                    if self.current_chapter_started:
                        # Skip separator after intro sections (TOC, Preface) - they're not content chapters
                        if ('table of contents' not in self.current_chapter_name.lower() and 
                            'contents' != self.current_chapter_name.lower() and 
                            'toc' != self.current_chapter_name.lower() and
                            'preface' != self.current_chapter_name.lower()):
                            separator_settings = self.config.get_chapter_separator()
                            if separator_settings and separator_settings.get('enabled') and separator_settings.get('position') == 'after':
                                self._add_chapter_separator(separator_settings)
                                if os.environ.get('WORD_FORMATTER_DEBUG', '0') == '1':
                                    print(f"DEBUG: Added chapter separator at end of '{self.current_chapter_name}' before starting '{heading_text}'")
                                else:
                                    print(f"Added separator at end of chapter: {self.current_chapter_name}")
                        elif os.environ.get('WORD_FORMATTER_DEBUG', '0') == '1':
                            print(f"DEBUG: Skipped separator after intro section: '{self.current_chapter_name}'")
                    
                    # Simple rule: Always add page break before chapters (H1 headings)
                    if len(self.output_doc.paragraphs) > 0:
                        self.output_doc.add_page_break()
                        if os.environ.get('WORD_FORMATTER_DEBUG', '0') == '1':
                            print(f"DEBUG: Added page break before chapter: '{heading_text}'")
                        
                        # Clear any special section flags since we're starting a new major section
                        if self.special_sections.get('dedication'):
                            self.special_sections['dedication'] = False
                        if self.special_sections.get('contents'):
                            self.special_sections['contents'] = False
                    
                    # Don't add separator before chapters - only after
                    # (The separator is added when starting the next chapter or at end of document)
                    
                    # Mark that we're in a chapter for end-of-chapter separator
                    self.current_chapter_started = True
                    self.current_chapter_name = heading_text
                    self.paragraphs_since_chapter = 0
                    self.current_chapter_elements = []
                # For other headings, shift down by 1 if we have chapters at level 2
                elif level >= 2:
                    word_heading_level = min(level + 1, 6)  # Cap at level 6
            else:
                # Fallback to hardcoded logic if no config
                # Check if this is a title first
                if 'title' in heading_text.lower():
                    use_title_style = True
                    word_heading_level = 0
                elif 'section' in heading_text.lower() or 'part' in heading_text.lower():
                    word_heading_level = 1
                elif level == 1 or 'chapter' in heading_text.lower():
                    word_heading_level = 2
                    if len(self.output_doc.paragraphs) > 0:
                        self.output_doc.add_page_break()
                elif level >= 2:
                    word_heading_level = min(level + 1, 6)
            
            # Check if this is a special section
            special_section = self._is_special_section(heading_text)
            
            # If this is the title (first H1), mark it as title special section
            if use_title_style:
                special_section = 'title'
            
            # Special section page breaks are handled in paragraph processing, not for headings
            # Headings get their page breaks from heading/chapter logic
            
            # Add the heading or title
            if use_title_style or special_section == 'title':
                # Use Title style instead of heading
                heading = self.output_doc.add_paragraph(heading_text)
                heading.style = 'Title'

                # Apply Title style overrides from configuration
                if self.config and hasattr(self.config, 'get_style_override'):
                    title_override = self.config.get_style_override('Title')
                    if title_override:
                        # Apply font properties to all runs
                        for run in heading.runs:
                            if 'font_name' in title_override and title_override['font_name'] is not None:
                                run.font.name = title_override['font_name']
                            if 'font_size' in title_override and title_override['font_size'] is not None:
                                run.font.size = Pt(title_override['font_size'])
                            if 'italic' in title_override and title_override['italic'] is not None:
                                run.italic = title_override['italic']
                            if 'bold' in title_override and title_override['bold'] is not None:
                                run.bold = title_override['bold']

                        # Apply paragraph formatting
                        if 'alignment' in title_override:
                            alignment_map = {
                                'left': WD_ALIGN_PARAGRAPH.LEFT,
                                'center': WD_ALIGN_PARAGRAPH.CENTER,
                                'right': WD_ALIGN_PARAGRAPH.RIGHT,
                                'justify': WD_ALIGN_PARAGRAPH.JUSTIFY
                            }
                            if title_override['alignment'] in alignment_map:
                                heading.alignment = alignment_map[title_override['alignment']]
                        if 'line_spacing' in title_override and title_override['line_spacing'] is not None:
                            heading.paragraph_format.line_spacing = title_override['line_spacing']

                # Mark as special section if detected
                if special_section:
                    self._handle_special_section_page_break(special_section)
            else:
                heading = self.output_doc.add_heading(heading_text, level=word_heading_level)
                # Check for other special sections
                if special_section:
                    self._handle_special_section_page_break(special_section)
            
            # Apply heading overrides from configuration if available
            if self.config and hasattr(self.config, 'get_heading_override'):
                override = self.config.get_heading_override(word_heading_level)
                if override:
                    # Apply alignment
                    if 'alignment' in override:
                        alignment_map = {
                            'left': WD_ALIGN_PARAGRAPH.LEFT,
                            'center': WD_ALIGN_PARAGRAPH.CENTER,
                            'right': WD_ALIGN_PARAGRAPH.RIGHT,
                            'justify': WD_ALIGN_PARAGRAPH.JUSTIFY
                        }
                        if override['alignment'] in alignment_map:
                            heading.alignment = alignment_map[override['alignment']]
                    
                    # Apply font properties to all runs
                    for run in heading.runs:
                        # First apply reference document font if available
                        if self.style_extractor and hasattr(self.style_extractor, 'styles') and 'headings' in self.style_extractor.styles:
                            heading_style_name = f'Heading {word_heading_level}'
                            if heading_style_name in self.style_extractor.styles['headings']:
                                ref_style = self.style_extractor.styles['headings'][heading_style_name]
                                if 'font' in ref_style and ref_style['font'].get('name'):
                                    run.font.name = ref_style['font']['name']

                        # Then apply overrides (only if explicitly set, not None)
                        if 'italic' in override and override['italic'] is not None:
                            run.italic = override['italic']
                        if 'bold' in override and override['bold'] is not None:
                            run.bold = override['bold']
                        if 'font_name' in override and override['font_name'] is not None:
                            run.font.name = override['font_name']
                        if 'font_size' in override and override['font_size'] is not None:
                            run.font.size = Pt(override['font_size'])
                    
                    # Note: Other formatting from template is preserved unless overridden
                else:
                    # No override - template styles are used
                    pass
            
            # Add page breaks after Title, Dedication, or Contents
            if self.config:
                # Check for Title page break
                if (hasattr(self.config, 'is_title_keyword') and self.config.is_title_keyword(heading_text) and 
                    self.config.should_add_page_break_after_title()):
                    # Add a page break after this paragraph
                    self.output_doc.add_page_break()
                
                # Check for Dedication page break
                elif (hasattr(self.config, 'is_dedication_keyword') and self.config.is_dedication_keyword(heading_text) and 
                      self.config.should_add_page_break_after_dedication()):
                    self.output_doc.add_page_break()
                
                # Check for Contents page break
                elif (hasattr(self.config, 'is_contents_keyword') and self.config.is_contents_keyword(heading_text) and 
                      self.config.should_add_page_break_after_contents()):
                    self.output_doc.add_page_break()
            
            # Check if this heading introduces a hierarchical list section
            if self.config:
                hl_settings = self.config.get_hierarchical_list_settings()
                if hl_settings.get('enabled') and hl_settings.get('detection_rules', {}).get('auto_detect_heading', True):
                    section_keywords = hl_settings.get('detection_rules', {}).get('section_keywords', [])
                    # Check if we're leaving a hierarchical list (new heading at same or higher level)
                    if self.in_hierarchical_list and self.current_list_heading_level is not None:
                        if word_heading_level <= self.current_list_heading_level:
                            self.in_hierarchical_list = False
                            self.current_list_heading_level = None
                    
                    # Check if this heading starts a hierarchical list
                    if any(keyword in heading_text.lower() for keyword in section_keywords):
                        self.in_hierarchical_list = True
                        self.current_list_heading_level = word_heading_level
            else:
                # Apply hardcoded formatting if no config (backward compatibility)
                if word_heading_level == 1 and ('section' in heading_text.lower() or 'part' in heading_text.lower()):
                    heading.paragraph_format.space_before = Pt(30)
                    heading.paragraph_format.space_after = Pt(24)
                    heading.paragraph_format.keep_with_next = True
                elif word_heading_level == 2 and ('chapter' in heading_text.lower() or level == 1):
                    heading.paragraph_format.space_before = Pt(24)
                    heading.paragraph_format.space_after = Pt(18)
                    heading.paragraph_format.keep_with_next = True
                elif word_heading_level == 3:
                    heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    for run in heading.runs:
                        run.italic = False
                elif word_heading_level == 4:
                    heading.alignment = WD_ALIGN_PARAGRAPH.LEFT
                    for run in heading.runs:
                        run.italic = True
        
        elif element.name == 'blockquote':
            # Handle blockquotes (Sanskrit quotes, etc.)
            paragraphs = element.find_all('p')
            
            if paragraphs:
                # Mark all child paragraphs as processed to avoid duplicates
                for p_elem in paragraphs:
                    processed_elements.add(id(p_elem))
                
                # Get blockquote settings from config
                if self.config:
                    bq_settings = self.config.get_blockquote_settings()
                    center_align = bq_settings.get('center_align', True)
                    remove_em_dashes = bq_settings.get('remove_em_dashes', True)
                    force_italic = bq_settings.get('italic', False)
                else:
                    center_align = True
                    remove_em_dashes = True
                    force_italic = False
                
                # Check if this is a chapter opening quote
                is_opening_quote = False
                if self.config and self.current_chapter_started and self.paragraphs_since_chapter <= 3:
                    blockquote_text = ' '.join([p.get_text().strip() for p in paragraphs])
                    if self._is_chapter_opening_content(blockquote_text, self.paragraphs_since_chapter):
                        is_opening_quote = True
                        if os.environ.get('WORD_FORMATTER_DEBUG', '0') == '1':
                            print(f"DEBUG: Chapter opening quote detected in blockquote: '{blockquote_text[:50]}...'")
                    elif os.environ.get('WORD_FORMATTER_DEBUG', '0') == '1':
                        print(f"DEBUG: Blockquote not detected as opening quote: '{blockquote_text[:50]}...', paragraphs_since_chapter={self.paragraphs_since_chapter}")
                
                # Create a single paragraph for the entire blockquote
                para = self.output_doc.add_paragraph()
                if center_align:
                    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                para.paragraph_format.space_before = Pt(12)
                para.paragraph_format.space_after = Pt(12)
                
                # Track if we're processing the transliteration (first non-source paragraph)
                transliteration_processed = False
                
                if os.environ.get('WORD_FORMATTER_DEBUG', '0') == '1':
                    print(f"DEBUG: Blockquote has {len(paragraphs)} paragraphs:")
                    for j, p in enumerate(paragraphs):
                        print(f"  Paragraph {j}: '{p.get_text().strip()}'")
                
                for i, p_elem in enumerate(paragraphs):
                    text = p_elem.get_text().strip()
                    
                    if text:
                        # Add line break between paragraphs (except for the first)
                        if i > 0:
                            para.add_run('\n')
                        
                        # Check if this looks like a source citation (last paragraph, often shorter)
                        # Look for patterns like author names, book titles, or short attributions
                        is_citation = False
                        if i == len(paragraphs) - 1:  # Last paragraph
                            # For multi-paragraph blockquotes, check if last paragraph is citation
                            if len(paragraphs) > 1:
                                if (len(text) < 100 and 
                                    (any(word[0].isupper() for word in text.split()) or 
                                     'Chapter' in text or 'Verse' in text or
                                     text.startswith('-') or text.startswith('—'))):
                                    is_citation = True
                            else:
                                # Single paragraph blockquote - handle quote and source separately
                                # Split by lines and check if the last line looks like a source
                                lines = text.split('\n')
                                if len(lines) > 1:
                                    last_line = lines[-1].strip()
                                    if (last_line.startswith('-') or last_line.startswith('—')) and len(last_line) < 100:
                                        # Process as separate quote and citation
                                        quote_lines = lines[:-1]
                                        quote_text = '\n'.join(quote_lines).strip()
                                        
                                        # Add the quote part (not italic by default)
                                        if is_opening_quote:
                                            run = para.add_run(quote_text)
                                            run.italic = True  # Quote part should be italic for opening quotes
                                        else:
                                            para.add_run(quote_text)
                                        
                                        # Add line break
                                        para.add_run('\n')
                                        
                                        # Add the citation part
                                        citation_text = last_line
                                        if not citation_text.startswith('—'):
                                            citation_text = '— ' + citation_text.lstrip('- ')
                                        run = para.add_run(citation_text)
                                        run.italic = True
                                        
                                        # Skip the normal processing for this paragraph
                                        continue
                        
                        if is_citation:
                            # Source citation - ensure it starts with em dash if it doesn't already
                            if not text.startswith('—') and not text.startswith('-'):
                                text = '— ' + text
                            elif text.startswith('-'):
                                # Replace hyphen with em dash
                                text = '—' + text[1:]
                            
                            # Apply reference style formatting for citations
                            run = para.add_run(text)
                            run.italic = True
                            # Use reference document's font if available
                            if hasattr(para, 'style') and para.style:
                                # Citation inherits from the paragraph style
                                pass
                        else:
                            # Check if the paragraph contains em/i tags
                            if p_elem.find('em') or p_elem.find('i'):
                                # Process inline elements to preserve italic formatting
                                for child in p_elem.children:
                                    if hasattr(child, 'name'):
                                        if child.name in ['em', 'i']:
                                            run = para.add_run(child.get_text())
                                            run.italic = True
                                            # Apply reference style for italic text in quotes
                                            if is_opening_quote:
                                                run.font.size = Pt(10)  # Smaller size for verses
                                        else:
                                            para.add_run(child.get_text())
                                    else:
                                        # Plain text
                                        text_content = str(child).strip()
                                        if text_content:
                                            para.add_run(text_content)
                            else:
                                # Plain text paragraph - for opening quotes, apply italic to entire text
                                if is_opening_quote and i == 0:
                                    # First paragraph of opening quote (usually the verse/quote itself)
                                    run = para.add_run(text)
                                    run.italic = True
                                else:
                                    para.add_run(text)
                
                # Apply opening quote formatting if needed
                if is_opening_quote:
                    # Apply Emphasis character style to all runs in the paragraph
                    for run in para.runs:
                        run.style = 'Emphasis'
                    if os.environ.get('WORD_FORMATTER_DEBUG', '0') == '1':
                        print(f"DEBUG: Applied Emphasis character style to chapter opening quote blockquote")

                # Apply italic formatting to all blockquote runs if configured
                if force_italic and not is_opening_quote:
                    for run in para.runs:
                        run.italic = True
                
                # Track this element for chapter end detection
                if self.current_chapter_started:
                    self.current_chapter_elements.append({
                        'type': 'blockquote',
                        'text': ' '.join([p.get_text().strip() for p in paragraphs]),
                        'paragraph': para
                    })
            
        elif element.name == 'p':
            text = element.get_text().strip()
            if os.environ.get('WORD_FORMATTER_DEBUG', '0') == '1':
                print(f"DEBUG: Processing paragraph element: '{text[:50]}...'")
            
            # Check for page break placeholder
            if '<!--PAGEBREAK-->' in str(element):
                if self.config and self.config.should_preserve_original_page_breaks():
                    self.output_doc.add_page_break()
                # Remove the placeholder from the text
                text = text.replace('<!--PAGEBREAK-->', '')
            
            # Check if this paragraph is a special section (like "Dedicated to")
            special_section = self._is_special_section(text)
            if os.environ.get('WORD_FORMATTER_DEBUG', '0') == '1':
                print(f"DEBUG: Checking paragraph for special section: '{text[:30]}...' -> {special_section}")
            if special_section and special_section == 'dedication':
                # Handle dedication as a special paragraph
                para = self.output_doc.add_paragraph()
                self._process_inline_elements(element, para)
                # Add page break after dedication paragraph
                self._handle_special_section_page_break(special_section)
                if os.environ.get('WORD_FORMATTER_DEBUG', '0') == '1':
                    print(f"DEBUG: Processed dedication paragraph (page break deferred until next major section)")
                return
            
            # Check if this is a chapter opening quote
            is_opening_quote = False
            if self.config and self.current_chapter_started and self.paragraphs_since_chapter <= 3:
                if self._is_chapter_opening_content(text, self.paragraphs_since_chapter):
                    is_opening_quote = True
                    if os.environ.get('WORD_FORMATTER_DEBUG', '0') == '1':
                        print(f"DEBUG: Chapter opening quote detected in paragraph: '{text[:50]}...'")
                elif os.environ.get('WORD_FORMATTER_DEBUG', '0') == '1':
                    print(f"DEBUG: Paragraph not detected as opening quote: '{text[:50]}...', paragraphs_since_chapter={self.paragraphs_since_chapter}")
            
            para = self.output_doc.add_paragraph()
            self._process_inline_elements(element, para)
            
            # Apply opening quote formatting if needed
            if is_opening_quote:
                # Apply Emphasis character style to all runs in the paragraph
                for run in para.runs:
                    run.style = 'Emphasis'
                if os.environ.get('WORD_FORMATTER_DEBUG', '0') == '1':
                    print(f"DEBUG: Applied Emphasis character style to chapter opening quote paragraph")
            
            # Check if this is part of a hierarchical list structure
            if self.config and self.in_hierarchical_list:
                hl_settings = self.config.get_hierarchical_list_settings()
                if hl_settings.get('enabled'):
                    # Check if this paragraph matches numbered item pattern
                    import re
                    numbered_pattern = hl_settings.get('detection_rules', {}).get('numbered_pattern', r'^\d+\.\s+[A-Z]')
                    if re.match(numbered_pattern, text):
                        # Apply Normal (Web) style if available
                        style_names = [style.name for style in self.output_doc.styles]
                        if os.environ.get('WORD_FORMATTER_DEBUG', '0') == '1':
                            print(f"DEBUG: Applying hierarchical list style to: '{text[:30]}...'")
                            print(f"DEBUG: Available styles: {[s for s in style_names if 'Normal' in s or 'Web' in s]}")
                        
                        target_style = hl_settings.get('numbered_item_style', 'Normal (Web)')
                        if target_style in style_names:
                            # Preserve existing font formatting before applying style
                            existing_fonts = []
                            for run in para.runs:
                                existing_fonts.append({
                                    'name': run.font.name,
                                    'size': run.font.size,
                                    'bold': run.font.bold,
                                    'italic': run.font.italic,
                                    'color': run.font.color.rgb if run.font.color and run.font.color.rgb else None
                                })
                            
                            # Apply the style
                            para.style = target_style
                            
                            # Restore font formatting that might have been overridden by style
                            for i, run in enumerate(para.runs):
                                if i < len(existing_fonts):
                                    font_info = existing_fonts[i]
                                    if font_info['name']:
                                        run.font.name = font_info['name']
                                    if font_info['size']:
                                        run.font.size = font_info['size']
                                    if font_info['bold'] is not None:
                                        run.font.bold = font_info['bold']
                                    if font_info['italic'] is not None:
                                        run.font.italic = font_info['italic']
                                    if font_info['color']:
                                        run.font.color.rgb = font_info['color']
                            
                            if os.environ.get('WORD_FORMATTER_DEBUG', '0') == '1':
                                print(f"DEBUG: Applied '{target_style}' style while preserving existing fonts")
                        else:
                            if os.environ.get('WORD_FORMATTER_DEBUG', '0') == '1':
                                print(f"DEBUG: Style '{target_style}' not found, using default")
                    elif text.strip().startswith('• ') or re.match(r'^\s+• ', text):
                        # This is a bullet point line in hierarchical list - add proper tab indentation
                        # Remove leading spaces and replace with tab indentation
                        bullet_text = text.lstrip()  # Remove leading spaces
                        if bullet_text.startswith('• '):
                            # Clear current paragraph content and add with tab indentation
                            para.clear()
                            para.add_run('\t' + bullet_text)  # Add tab + bullet point text
                            
                            # Apply the bullet point style
                            bullet_style = hl_settings.get('bullet_point_style', 'List Paragraph')
                            style_names = [style.name for style in self.output_doc.styles]
                            if bullet_style in style_names:
                                para.style = bullet_style

                            # Apply Normal style overrides (including Aptos Light font) to list items
                            if self.config and hasattr(self.config, 'get_style_override'):
                                normal_override = self.config.get_style_override('Normal')
                                if normal_override and para.runs:
                                    for run in para.runs:
                                        if 'font_name' in normal_override and normal_override['font_name'] is not None:
                                            run.font.name = normal_override['font_name']
                                        if 'font_size' in normal_override and normal_override['font_size'] is not None:
                                            run.font.size = Pt(normal_override['font_size'])

                            if os.environ.get('WORD_FORMATTER_DEBUG', '0') == '1':
                                print(f"DEBUG: Applied tab indentation to bullet point: '{bullet_text[:30]}...'")
                    else:
                        # Check if we've left the hierarchical list section
                        # (e.g., if we encounter a regular paragraph that doesn't match the pattern)
                        if not text.strip().startswith('-') and not text.strip().startswith('•'):
                            self.in_hierarchical_list = False
                            self.current_list_heading_level = None
            
            # Track this element for chapter end detection
            if self.current_chapter_started and text:
                self.current_chapter_elements.append({
                    'type': 'paragraph',
                    'text': text,
                    'paragraph': para
                })
            
        elif element.name in ['ul', 'ol']:
            # Check if we should use special list formatting
            use_list_paragraph_style = False
            if self.config and self.in_hierarchical_list:
                hl_settings = self.config.get_hierarchical_list_settings()
                if hl_settings.get('enabled'):
                    use_list_paragraph_style = True
            
            for li in element.find_all('li', recursive=False):
                para = self.output_doc.add_paragraph(li.get_text().strip())
                # Try to use list styles if available
                try:
                    if use_list_paragraph_style and 'List Paragraph' in [style.name for style in self.output_doc.styles]:
                        para.style = 'List Paragraph'
                    elif element.name == 'ul' and 'List Bullet' in [style.name for style in self.output_doc.styles]:
                        para.style = 'List Bullet'
                    elif element.name == 'ol' and 'List Number' in [style.name for style in self.output_doc.styles]:
                        para.style = 'List Number'
                    else:
                        # Fallback to normal style with manual formatting
                        para.style = 'Normal'
                        if element.name == 'ul':
                            para.text = '• ' + para.text
                        else:
                            # Get the list item number
                            list_items = element.find_all('li', recursive=False)
                            item_index = list_items.index(li) + 1
                            para.text = f'{item_index}. ' + para.text
                except:
                    # If any style error, use normal style
                    para.style = 'Normal'
                
        elif element.name == 'table':
            self._process_table(element)
            
        elif element.name == 'pre':
            # Code block
            para = self.output_doc.add_paragraph(element.get_text())
            para.style = 'Quote'  # Use Quote style for code blocks
    
    def _process_inline_elements(self, element, paragraph):
        """Process inline formatting like bold, italic, links."""
        for child in element.children:
            if hasattr(child, 'name'):
                if child.name == 'strong' or child.name == 'b':
                    run = paragraph.add_run(child.get_text())
                    run.bold = True
                elif child.name == 'em' or child.name == 'i':
                    run = paragraph.add_run(child.get_text())
                    run.italic = True
                elif child.name == 'code':
                    run = paragraph.add_run(child.get_text())
                    run.font.name = 'Courier New'
                elif child.name == 'a':
                    run = paragraph.add_run(child.get_text())
                    run.font.color.rgb = RGBColor(0, 0, 255)
                    run.underline = True
                else:
                    paragraph.add_run(child.get_text())
            else:
                # Plain text
                paragraph.add_run(str(child))
    
    def _process_table(self, table_element):
        """Process HTML table to Word table."""
        rows = table_element.find_all('tr')
        if not rows:
            return
        
        # Count columns
        cols = len(rows[0].find_all(['td', 'th']))
        
        # Create table
        table = self.output_doc.add_table(rows=len(rows), cols=cols)
        table.style = 'Table Grid'
        
        # Populate table
        for i, row in enumerate(rows):
            cells = row.find_all(['td', 'th'])
            for j, cell in enumerate(cells):
                table.rows[i].cells[j].text = cell.get_text().strip()
    
    def _process_word_document(self, file_path: Path):
        """Process existing Word document with full formatting support."""
        # Resolve any aliases
        resolved_path = resolve_path(str(file_path))
        source_doc = Document(resolved_path)
        
        # Process each paragraph in order
        for para in source_doc.paragraphs:
            # Check if we should preserve original page breaks
            if self.config and self.config.should_preserve_original_page_breaks():
                # Check if this paragraph has a page break before it
                if para.paragraph_format.page_break_before:
                    self.output_doc.add_page_break()
            
            # Skip empty paragraphs but check for page break runs
            if not para.text.strip():
                # Check if this paragraph contains a page break character
                for run in para.runs:
                    # Check for page break in run's XML
                    if run._element.xml.find('<w:br w:type="page"/>') != -1:
                        if self.config and self.config.should_preserve_original_page_breaks():
                            self.output_doc.add_page_break()
                            break
                self.output_doc.add_paragraph()
                continue
            
            # Check if it's a heading
            if para.style and para.style.name and 'Heading' in para.style.name:
                # Extract heading level from style name (e.g., "Heading 1" -> 1)
                try:
                    level = int(para.style.name.split()[-1])
                except:
                    level = 1
                
                heading_text = para.text.strip()
                
                # Determine the Word heading level using configuration
                word_heading_level = level
                use_title_style = False
                
                if self.config:
                    # Check if this is a title
                    if hasattr(self.config, 'is_title_keyword') and self.config.is_title_keyword(heading_text):
                        use_title_style = True
                        word_heading_level = 0
                    # Check if this is a section (use Heading 1)
                    elif self.config.is_section_keyword(heading_text):
                        word_heading_level = 1
                        # Check for page break setting
                        if self.config.should_apply_page_break("section") and len(self.output_doc.paragraphs) > 0:
                            self.output_doc.add_page_break()
                    # Check if this is a chapter (always use Heading 2)
                    elif level == 1 or self.config.is_chapter_keyword(heading_text):
                        word_heading_level = 2
                        
                        # Add chapter separator if configured
                        separator_settings = self.config.get_chapter_separator()
                        if separator_settings.get('enabled') and len(self.output_doc.paragraphs) > 0:
                            # Add page break first if configured
                            if self.config.should_apply_page_break("chapter"):
                                self.output_doc.add_page_break()
                            
                            # Add separator before chapter
                            if separator_settings.get('position') == 'before':
                                sep_para = self.output_doc.add_paragraph()
                                sep_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                                sep_para.paragraph_format.space_before = Pt(separator_settings.get('spacing_before', 12))
                                sep_para.paragraph_format.space_after = Pt(separator_settings.get('spacing_after', 12))
                                
                                # Use image if available, otherwise use text symbol
                                if separator_settings.get('use_image') and separator_settings.get('image_data'):
                                    import io
                                    image_data = separator_settings['image_data']
                                    width = separator_settings.get('image_width', 0.5)
                                    height = separator_settings.get('image_height', 0.5)
                                    sep_para.add_run().add_picture(io.BytesIO(image_data), 
                                                                 width=Inches(width), 
                                                                 height=Inches(height))
                                else:
                                    # Use text symbol
                                    symbol = separator_settings.get('symbol', '❦')
                                    run = sep_para.add_run(symbol)
                                    run.font.size = Pt(separator_settings.get('font_size', 14))
                        else:
                            # Just page break if no separator
                            if self.config.should_apply_page_break("chapter") and len(self.output_doc.paragraphs) > 0:
                                self.output_doc.add_page_break()
                    # For other headings, shift down by 1 if we have chapters at level 2
                    elif level >= 2:
                        word_heading_level = min(level + 1, 6)
                else:
                    # Fallback logic without config
                    if 'section' in heading_text.lower() or 'part' in heading_text.lower():
                        word_heading_level = 1
                    elif level == 1 or 'chapter' in heading_text.lower():
                        word_heading_level = 2
                        if len(self.output_doc.paragraphs) > 0:
                            self.output_doc.add_page_break()
                    elif level >= 2:
                        word_heading_level = min(level + 1, 6)
                
                # Add the heading or title
                if use_title_style:
                    # Use Title style instead of heading
                    heading = self.output_doc.add_paragraph(heading_text)
                    heading.style = 'Title'

                    # Apply Title style overrides from configuration
                    if self.config and hasattr(self.config, 'get_style_override'):
                        title_override = self.config.get_style_override('Title')
                        if title_override:
                            # Apply font properties to all runs
                            for run in heading.runs:
                                if 'font_name' in title_override and title_override['font_name'] is not None:
                                    run.font.name = title_override['font_name']
                                if 'font_size' in title_override and title_override['font_size'] is not None:
                                    run.font.size = Pt(title_override['font_size'])
                                if 'italic' in title_override and title_override['italic'] is not None:
                                    run.italic = title_override['italic']
                                if 'bold' in title_override and title_override['bold'] is not None:
                                    run.bold = title_override['bold']

                            # Apply paragraph formatting
                            if 'alignment' in title_override:
                                alignment_map = {
                                    'left': WD_ALIGN_PARAGRAPH.LEFT,
                                    'center': WD_ALIGN_PARAGRAPH.CENTER,
                                    'right': WD_ALIGN_PARAGRAPH.RIGHT,
                                    'justify': WD_ALIGN_PARAGRAPH.JUSTIFY
                                }
                                if title_override['alignment'] in alignment_map:
                                    heading.alignment = alignment_map[title_override['alignment']]
                            if 'line_spacing' in title_override and title_override['line_spacing'] is not None:
                                heading.paragraph_format.line_spacing = title_override['line_spacing']
                else:
                    heading = self.output_doc.add_heading(heading_text, level=word_heading_level)
                
                # Apply heading overrides from configuration if available
                if self.config and hasattr(self.config, 'get_heading_override'):
                    override = self.config.get_heading_override(word_heading_level)
                    if override:
                        # Apply alignment
                        if 'alignment' in override:
                            alignment_map = {
                                'left': WD_ALIGN_PARAGRAPH.LEFT,
                                'center': WD_ALIGN_PARAGRAPH.CENTER,
                                'right': WD_ALIGN_PARAGRAPH.RIGHT,
                                'justify': WD_ALIGN_PARAGRAPH.JUSTIFY
                            }
                            if override['alignment'] in alignment_map:
                                heading.alignment = alignment_map[override['alignment']]

                        # Apply font properties to all runs
                        for run in heading.runs:
                            if 'italic' in override and override['italic'] is not None:
                                run.italic = override['italic']
                            if 'bold' in override and override['bold'] is not None:
                                run.bold = override['bold']
                            if 'font_name' in override and override['font_name'] is not None:
                                run.font.name = override['font_name']
                            if 'font_size' in override and override['font_size'] is not None:
                                run.font.size = Pt(override['font_size'])
            
            # Process lists
            elif para.style and para.style.name and ('List' in para.style.name or 'Bullet' in para.style.name):
                new_para = self.output_doc.add_paragraph(para.text)
                # Apply list style if available
                try:
                    if 'Bullet' in para.style.name and 'List Bullet' in [style.name for style in self.output_doc.styles]:
                        new_para.style = 'List Bullet'
                    elif 'List Number' in [style.name for style in self.output_doc.styles]:
                        new_para.style = 'List Number'
                    else:
                        new_para.style = 'Normal'
                except:
                    new_para.style = 'Normal'
                # Copy run formatting
                if para.runs and new_para.runs:
                    for i, run in enumerate(para.runs):
                        if i < len(new_para.runs):
                            new_para.runs[i].bold = run.bold
                            new_para.runs[i].italic = run.italic
                            new_para.runs[i].underline = run.underline
                            if run.font.color and run.font.color.rgb:
                                new_para.runs[i].font.color.rgb = run.font.color.rgb
            
            # Process regular paragraphs
            else:
                new_para = self.output_doc.add_paragraph()
                
                # Check if this is a blockquote (indented paragraph with specific formatting)
                is_blockquote = (para.paragraph_format.left_indent and 
                               para.paragraph_format.left_indent > Pt(0) and
                               para.alignment == WD_ALIGN_PARAGRAPH.CENTER)
                
                if is_blockquote and self.config:
                    # Apply blockquote settings
                    bq_settings = self.config.get_blockquote_settings()
                    if bq_settings.get('center_align', True):
                        new_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    new_para.paragraph_format.space_before = Pt(12)
                    new_para.paragraph_format.space_after = Pt(12)
                    force_italic_bq = bq_settings.get('italic', False)
                else:
                    force_italic_bq = False

                # Copy paragraph alignment if not a blockquote
                if para.alignment is not None and not is_blockquote:
                    new_para.alignment = para.alignment
                
                # Copy runs with formatting
                for run in para.runs:
                    text = run.text
                    # Handle blockquote citation processing
                    if is_blockquote and self.config:
                        bq_settings = self.config.get_blockquote_settings()
                        if bq_settings.get('remove_em_dashes', True) and text.startswith('—'):
                            text = text[1:].strip()
                    
                    new_run = new_para.add_run(text)
                    new_run.bold = run.bold
                    new_run.italic = run.italic if not force_italic_bq else True
                    new_run.underline = run.underline
                    if run.font.color and run.font.color.rgb:
                        new_run.font.color.rgb = run.font.color.rgb
                    if run.font.size:
                        new_run.font.size = run.font.size
                    if run.font.name:
                        new_run.font.name = run.font.name
        
        # Process tables
        for table in source_doc.tables:
            # Create new table with same dimensions
            new_table = self.output_doc.add_table(rows=len(table.rows), cols=len(table.columns))
            new_table.style = table.style
            
            # Copy cell contents
            for i, row in enumerate(table.rows):
                for j, cell in enumerate(row.cells):
                    new_table.rows[i].cells[j].text = cell.text
                    # Copy cell formatting if needed
                    if cell.paragraphs:
                        for k, para in enumerate(cell.paragraphs):
                            if k < len(new_table.rows[i].cells[j].paragraphs):
                                new_cell_para = new_table.rows[i].cells[j].paragraphs[k]
                                if para.alignment:
                                    new_cell_para.alignment = para.alignment
        
        # Process images and other inline shapes
        self._process_document_images(source_doc)
    
    def _process_document_images(self, source_doc):
        """Process and copy images from source document."""
        # For now, we'll skip complex image processing to avoid XML parsing issues
        # This is a placeholder for future enhancement
        # Images in Word documents require complex XML manipulation that can vary
        # between python-docx versions
        pass
    
    def _apply_reference_styles(self):
        """Apply styles from reference document to output document."""
        # Apply reference styles if style extractor is available
        if self.use_config and self.style_extractor:
            styles = self.style_extractor.styles
        else:
            styles = None


        # Apply document-level styles from reference if available
        if styles and 'document' in styles and 'margins' in styles['document']:
            section = self.output_doc.sections[0]
            margins = styles['document']['margins']
            if margins['top']:
                section.top_margin = margins['top']
            if margins['bottom']:
                section.bottom_margin = margins['bottom']
            if margins['left']:
                section.left_margin = margins['left']
            if margins['right']:
                section.right_margin = margins['right']

        # Apply normal style from reference if available
        if styles and 'normal' in styles and styles['normal']:
            normal_style = styles['normal']
            for para in self.output_doc.paragraphs:
                if para.style.name == 'Normal':
                    self._apply_paragraph_style(para, normal_style)

        # ALWAYS apply Normal style overrides from configuration if config exists
        if self.config and hasattr(self.config, 'get_style_override'):
            normal_override = self.config.get_style_override('Normal')
            if normal_override:
                for para in self.output_doc.paragraphs:
                    if para.style.name == 'Normal':
                        # Apply font properties to all runs
                        for run in para.runs:
                            if 'font_name' in normal_override and normal_override['font_name'] is not None:
                                run.font.name = normal_override['font_name']
                            if 'font_size' in normal_override and normal_override['font_size'] is not None:
                                run.font.size = Pt(normal_override['font_size'])
                            if 'italic' in normal_override and normal_override['italic'] is not None:
                                run.italic = normal_override['italic']
                            if 'bold' in normal_override and normal_override['bold'] is not None:
                                run.bold = normal_override['bold']

                        # Apply paragraph formatting
                        if 'line_spacing' in normal_override and normal_override['line_spacing'] is not None:
                            para.paragraph_format.line_spacing = normal_override['line_spacing']
                        if 'alignment' in normal_override:
                            alignment_map = {
                                'left': WD_ALIGN_PARAGRAPH.LEFT,
                                'center': WD_ALIGN_PARAGRAPH.CENTER,
                                'right': WD_ALIGN_PARAGRAPH.RIGHT,
                                'justify': WD_ALIGN_PARAGRAPH.JUSTIFY
                            }
                            if normal_override['alignment'] in alignment_map:
                                para.alignment = alignment_map[normal_override['alignment']]
        
        # Apply heading styles from reference if available
        if styles and 'headings' in styles:
            for para in self.output_doc.paragraphs:
                if para.style and 'Heading' in para.style.name:
                    if para.style.name in styles['headings']:
                        self._apply_paragraph_style(para, styles['headings'][para.style.name])
                # Also apply Title style if present
                elif para.style and para.style.name == 'Title' and 'Title' in styles['headings']:
                    self._apply_paragraph_style(para, styles['headings']['Title'])
    
    def _apply_paragraph_style(self, paragraph: Paragraph, style: Dict):
        """Apply style dictionary to a paragraph."""
        # Apply paragraph formatting
        if 'paragraph' in style:
            para_format = style['paragraph']
            if para_format.get('alignment'):
                paragraph.alignment = para_format['alignment']
            if para_format.get('space_before'):
                paragraph.paragraph_format.space_before = para_format['space_before']
            if para_format.get('space_after'):
                paragraph.paragraph_format.space_after = para_format['space_after']
            if para_format.get('line_spacing'):
                paragraph.paragraph_format.line_spacing = para_format['line_spacing']
            if para_format.get('left_indent'):
                paragraph.paragraph_format.left_indent = para_format['left_indent']
            if para_format.get('right_indent'):
                paragraph.paragraph_format.right_indent = para_format['right_indent']
            if para_format.get('first_line_indent'):
                paragraph.paragraph_format.first_line_indent = para_format['first_line_indent']
            if para_format.get('keep_with_next'):
                paragraph.paragraph_format.keep_with_next = para_format['keep_with_next']
            if para_format.get('page_break_before'):
                paragraph.paragraph_format.page_break_before = para_format['page_break_before']
        
        # Apply font formatting to all runs
        if 'font' in style and paragraph.runs:
            font_style = style['font']
            for run in paragraph.runs:
                if font_style.get('name'):
                    run.font.name = font_style['name']
                if font_style.get('size'):
                    run.font.size = font_style['size']
                if font_style.get('bold') is not None:
                    run.font.bold = font_style['bold']
                if font_style.get('italic') is not None:
                    run.font.italic = font_style['italic']
                if font_style.get('color'):
                    run.font.color.rgb = font_style['color']


def main():
    """Main entry point for the document converter."""
    parser = argparse.ArgumentParser(
        description='Convert text, markdown, or Word documents to match a reference Word document style.',
        epilog='Example: python document_converter.py --input document.txt --reference template.docx --output result.docx'
    )
    parser.add_argument('--input', '-i', dest='input_file',
                        help='Path to input file (txt, md, or docx)')
    parser.add_argument('--reference', '-r', required=True, dest='reference_file',
                        help='Path to reference Word document for styling')
    parser.add_argument('--output', '-o', dest='output',
                        help='Output file path (default: input_file_formatted.docx)')
    parser.add_argument('--config', '-c', dest='config_file',
                        help='Path to configuration file (default: formatter_config.json)')
    parser.add_argument('--export-styles', dest='export_styles',
                        help='Export all styles from reference document to JSON file')
    
    args = parser.parse_args()
    
    # Validate inputs
    reference_path = Path(args.reference_file)
    
    # If exporting styles, we don't need an input file
    if args.export_styles:
        if not reference_path.exists():
            print(f"Error: Reference file does not exist: {reference_path}")
            sys.exit(1)
    else:
        # For normal conversion, input file is required
        if not args.input_file:
            print("Error: --input is required when not exporting styles")
            sys.exit(1)
            
        input_path = Path(args.input_file)
        if not input_path.exists():
            print(f"Error: Input file does not exist: {input_path}")
            sys.exit(1)
    
    if not reference_path.exists():
        print(f"Error: Reference file does not exist: {reference_path}")
        sys.exit(1)
    
    if reference_path.suffix.lower() not in ['.docx', '.doc']:
        print(f"Error: Reference file must be a Word document (.docx or .doc)")
        sys.exit(1)
    
    # Determine output path (only needed for conversion, not style export)
    if not args.export_styles:
        if args.output:
            output_path = Path(args.output)
        else:
            output_path = input_path.parent / f"{input_path.stem}_formatted.docx"
    
    # Create converter
    try:
        # First, extract styles if requested
        if args.export_styles:
            print(f"Extracting styles from: {reference_path}")
            style_extractor = StyleExtractor(reference_path)
            style_extractor._log_extracted_styles()  # Show debug output
            style_extractor.export_styles_to_file(args.export_styles)
            print(f"\nStyle extraction complete!")
            return
        
        # Otherwise, do normal conversion
        converter = DocumentConverter(reference_path, config_path=args.config_file)
        converter.convert(str(input_path), str(output_path))
    except Exception as e:
        print(f"Error during conversion: {e}")
        sys.exit(1)


if __name__ == "__main__":
    main()