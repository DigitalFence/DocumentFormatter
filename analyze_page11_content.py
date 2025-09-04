#!/usr/bin/env python3
"""Analyze page 11 content in reference document."""

from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from alias_resolver import resolve_path

# Resolve the alias
doc_path = resolve_path('referenceformat.docx')
doc = Document(doc_path)

print('=== ANALYZING CONTENT AROUND PAGE 11 ===')
print()

# With 170 paragraphs total, page 11 would be around paragraphs 140-160
# Let's look at a wider range to be sure
start = 130
end = 170

print(f'Examining paragraphs {start} to {end}')
print(f'Total document paragraphs: {len(doc.paragraphs)}')
print()

special_formatting_found = []

for i in range(start, min(end, len(doc.paragraphs))):
    para = doc.paragraphs[i]
    
    # Skip empty paragraphs
    if not para.text.strip():
        continue
    
    # Check for any special characteristics
    special = []
    
    # Alignment
    if para.alignment == WD_ALIGN_PARAGRAPH.CENTER:
        special.append('CENTER')
    elif para.alignment == WD_ALIGN_PARAGRAPH.RIGHT:
        special.append('RIGHT')
    
    # Style
    if para.style and para.style.name != 'Normal':
        special.append(f'STYLE={para.style.name}')
    
    # Indentation
    pf = para.paragraph_format
    if pf.left_indent and pf.left_indent.pt > 0:
        special.append(f'LEFT_INDENT={pf.left_indent.pt}pt')
    if pf.right_indent and pf.right_indent.pt > 0:
        special.append(f'RIGHT_INDENT={pf.right_indent.pt}pt')
    
    # Font characteristics
    if para.runs:
        fonts = set()
        sizes = set()
        has_italic = False
        has_bold = False
        
        for run in para.runs:
            if run.text.strip():
                if run.font.name:
                    fonts.add(run.font.name)
                if run.font.size:
                    sizes.add(run.font.size.pt)
                if run.italic:
                    has_italic = True
                if run.bold:
                    has_bold = True
        
        if fonts and 'Aptos Light' not in fonts:
            special.append(f'FONT={list(fonts)}')
        if sizes:
            special.append(f'SIZE={list(sizes)}')
        if has_italic:
            special.append('ITALIC')
        if has_bold:
            special.append('BOLD')
    
    # Print paragraph info
    print(f'Para {i}:')
    text_preview = para.text[:80] + '...' if len(para.text) > 80 else para.text
    print(f'  Text: {text_preview}')
    if special:
        print(f'  Special: {" | ".join(special)}')
        special_formatting_found.append({
            'index': i,
            'text': text_preview,
            'formatting': special
        })
    print()

print('\n=== SUMMARY OF SPECIAL FORMATTING ===')
print(f'Found {len(special_formatting_found)} paragraphs with special formatting')

# Group by formatting type
formatting_types = {}
for item in special_formatting_found:
    for fmt in item['formatting']:
        if fmt not in formatting_types:
            formatting_types[fmt] = 0
        formatting_types[fmt] += 1

print('\nFormatting types found:')
for fmt_type, count in sorted(formatting_types.items(), key=lambda x: x[1], reverse=True):
    print(f'  {fmt_type}: {count} occurrences')